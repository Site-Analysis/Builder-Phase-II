"""Generate SAT_Deployment_Plan.docx from the approved deployment plan.

One-off documentation builder. Mirrors the plan at
.claude/plans/hey-claude-use-this-peppy-russell.md plus the API inventory and
building-precision sections. Run: python scripts/gen_deployment_doc.py
"""

from __future__ import annotations

import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# ── palette (SAT brand) ──────────────────────────────────────────────────────
BRAND = RGBColor(0x2E, 0x7D, 0x6F)      # primary teal
INK = RGBColor(0x1C, 0x24, 0x20)        # near-black
MUTED = RGBColor(0x5A, 0x6A, 0x66)
MONO_BG = "EAF2F1"                       # secondary tint


def shade(cell, hex_fill: str) -> None:
    tc = cell._tc.get_or_add_tcPr()
    shd = tc.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tc.append(shd)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = BRAND
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pbdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pbdr.makeelement(qn("w:bottom"), {})
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "2E7D6F")
        pbdr.append(bottom)
        pPr.append(pbdr)
    else:
        run.font.size = Pt(12.5)
        run.font.color.rgb = INK
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)


def body(doc: Document, text: str, bold=False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(4)


def bullets(doc: Document, items, numbered=False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for it in items:
        p = doc.add_paragraph(style=style)
        r = p.add_run(it)
        r.font.size = Pt(10.5)
        r.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(1)


def mono(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(6)
    for line in text.split("\n"):
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = INK
    shade_paragraph(p, MONO_BG)


def shade_paragraph(p, hex_fill: str) -> None:
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


def table(doc: Document, headers, rows) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], "2E7D6F")
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if ri % 2 == 1:
                shade(cells[i], "F4F8F7")
            para = cells[i].paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── document ─────────────────────────────────────────────────────────────────
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

# title page
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = tp.add_run("SAT — Deployment Plan")
tr.bold = True
tr.font.size = Pt(26)
tr.font.color.rgb = BRAND
st = doc.add_paragraph()
sr = st.add_run("Beta · AWS EC2 + Vercel · Full Architecture, API Inventory & Gates")
sr.font.size = Pt(12)
sr.font.color.rgb = MUTED

meta = doc.add_paragraph()
meta.paragraph_format.space_before = Pt(8)
for label, val in [
    ("Date", datetime.date.today().isoformat()),
    ("Branch", "ux/phase-1-research"),
    ("Author", "Tanmay C J"),
    ("Status", "Phase 10 — Beta Readiness: GO"),
    ("Scope", "10 FastAPI services + Next.js 16 frontend + Supabase"),
]:
    rr = meta.add_run(f"{label}:  ")
    rr.bold = True
    rr.font.size = Pt(10)
    rr.font.color.rgb = INK
    vv = meta.add_run(f"{val}\n")
    vv.font.size = Pt(10)
    vv.font.color.rgb = MUTED

# 2. Executive summary
heading(doc, "1. Executive Summary")
body(doc, "SAT is Beta-ready (Phase 10 GO). Goal: get the full stack live for stakeholders and a "
          "few users at low cost. Frontend is Next.js 16; backend is 10 FastAPI microservices "
          "(ports 8000-8009), already containerized with per-service Dockerfiles and a root "
          "docker-compose.yml. Supabase (auth/persistence) is managed/external. Only 2 services "
          "are heavy (temperature, sunpath); the other 8 are light proxies to external APIs.")
body(doc, "Locked decisions:", bold=True)
bullets(doc, [
    "Budget ~$20-40/month.",
    "Frontend on Vercel free tier (native Next.js 16).",
    "Backend on a single AWS EC2 box (t3.medium) via docker-compose.",
    "Maturity target: beta / low-traffic.",
])

# 3. Phase 1 — visuals
heading(doc, "2. How the Visuals Work Today (Phase 1)")
table(doc, ["Concern", "Tech stack + theory"], [
    ["3D buildings", "maplibre-gl basemap + three.js (MapLibre custom layer, @react-three/fiber + drei). "
                     "Footprints from MapLibre vector-tile 'building' layer; heights from OSM tags; "
                     "THREE.Shape -> ExtrudeGeometry -> merged mesh. components/three/Scene3D.tsx. Needs MapTiler vector-tile key."],
    ["Shadows 3D", "Real-time shadow mapping. DirectionalLight set from sun az/el (sunpath SolarPoint[]), "
                   "PCF 2048^2 shadow map, orthographic frustum on site, ShadowMaterial ground plane."],
    ["Shadows 2D", "Vector projection, no raster. lib/geo.ts shadowPolygon() sweeps footprint by "
                   "height/tan(elev) along anti-sun bearing -> convex hull -> Leaflet Polygon."],
    ["Amenities", "OpenStreetMap Overpass API, proxied via geo service /geo/amenities (8005). "
                  "7 categories, 2km default. Rendered as Markers + recharts bar chart."],
    ["Analysis visuals", "2D: react-leaflet + leaflet, CARTO Voyager tiles (free). 3D: maplibre-gl + three. "
                         "All overlays client-side vector (Circle/Polygon/Polyline/Marker). zustand state, "
                         "recharts, html2canvas + jspdf export."],
])
body(doc, "Heavy services: temperature (8000 - netCDF4/xarray/scipy/imdlib + large IMD .grd data mount) "
          "and sunpath (8001 - earthengine-api/pvlib/matplotlib/shapely/pyproj + GEE service account). "
          "Other 8 are light httpx -> external-API proxies (Overpass/Open-Meteo/GEE).")

# 4. Building precision
heading(doc, "3. How Buildings Are Precision-Analyzed Before Render")
body(doc, "Two client-side pipelines, both OSM-sourced. No GIS engine (no PostGIS/turf/shapely) - "
          "hand-rolled TypeScript geometry + three.js.")
heading(doc, "2D path (Leaflet) - lib/osm.ts fetchBuildings()", level=2)
bullets(doc, [
    "Overpass way[\"building\"](around:radius) -> out geom.",
    "Reject ways with <3 vertices.",
    "Height = tag-driven: height stripped to numeric, else building:levels x 3, else null (parseHeight). Not measured.",
    "Sort nearest-first, cap 400, cache by 4-decimal coord key.",
], numbered=True)
heading(doc, "3D path (MapLibre + three.js) - Scene3D.tsx buildingsFromMap()", level=2)
bullets(doc, [
    "Footprints read from MapLibre's own in-browser vector tiles - no network, shadows align with visible extrusions.",
    "Outer ring only (Polygon/MultiPolygon).",
    "Dedupe tile-split features (Set by id/osm_id/coord).",
    "Height = render_height ?? height ?? 0 || null; default 10m at extrude.",
    "Reject <3 verts; try/catch skips malformed polygons.",
    "lat/lng -> metres via equirectangular tangent-plane (cosLat correction, exact at site scale).",
    "THREE.Shape -> ExtrudeGeometry(depth=height) -> mergeGeometries.",
], numbered=True)
body(doc, "Tools: OpenStreetMap (Overpass for 2D, MapLibre vector tiles for 3D) + custom TS geo "
          "(lib/geo.ts: Andrew monotone-chain convex hull, ray-cast point-in-polygon, equirectangular "
          "projection) + three.js extrusion.")
body(doc, "Precision limits:", bold=True)
bullets(doc, [
    "Heights = OSM tags or 10m default (not surveyed).",
    "Equirectangular projection - accurate only at small site radius.",
    "Footprint accuracy = whatever OSM has; no satellite/LiDAR refinement.",
    "400-building cap; outer ring only (holes dropped).",
    "Not a rigorous GIS topology-validation pipeline.",
])

# 5. Heavy vs light
heading(doc, "4. Heavy vs Light Services")
table(doc, ["Service", "Port", "Weight", "Why"], [
    ["temperature", "8000", "HEAVY", "netCDF4, xarray, scipy, imdlib (native wheels) + large IMD .grd data mount"],
    ["sunpath", "8001", "HEAVY", "earthengine-api, pvlib, matplotlib, shapely, pyproj + GEE service account"],
    ["flood", "8002", "light", "httpx + pydantic; external API only"],
    ["wind", "8003", "light", "httpx + pydantic; Open-Meteo"],
    ["rainfall", "8004", "light", "httpx + pydantic; Open-Meteo / GEE optional"],
    ["geo", "8005", "light", "httpx; Overpass + 266KB GeoJSON"],
    ["planning", "8006", "light", "httpx + pydantic; no external deps"],
    ["infrastructure", "8007", "light", "httpx; Overpass"],
    ["future-infra", "8008", "light", "httpx; small local config JSON"],
    ["land-records", "8009", "light", "httpx + pydantic; no external deps"],
])

# 6. Decisions
heading(doc, "5. Locked Decisions")
table(doc, ["Decision", "Choice", "Rationale"], [
    ["Budget", "~$20-40/mo", "Run real compute; near-$0 first 6mo on AWS credits"],
    ["Frontend host", "Vercel free", "Native Next.js 16, best DX/perf, generous free tier"],
    ["Backend compute", "Single EC2 + docker-compose", "Simplest lift-and-shift; one box for beta"],
    ["Maturity", "Beta / low-traffic", "Matches Phase 10 Beta Readiness: GO"],
])

# 7. Architecture
heading(doc, "6. Target Architecture")
mono(doc,
     "[Vercel] Next.js 16 (free tier)\n"
     "   |- NEXT_PUBLIC_*_API_URL = https://api.<domain>\n"
     "            |  HTTPS, CORS locked to Vercel origin\n"
     "            v\n"
     "[AWS ap-south-1] 1x EC2 t3.medium (4GB) + Elastic IP\n"
     "   |- Caddy reverse proxy (auto-TLS, path-prefix routing)\n"
     "        |- /weather/*    -> temperature:8000   (heavy)\n"
     "        |- /sunpath/*    -> sunpath:8001        (heavy)\n"
     "        |- /flood/* /wind/* /rainfall/* /geo/*\n"
     "        |- /planning/* /infrastructure/*\n"
     "        |- /future-infra/* /land-records/*      (8 light)\n"
     "   docker-compose up -d  (all 10 services + Caddy)\n"
     "   EBS gp3 30GB: docker images + IMD data + gee-sa.json\n"
     "[Supabase] managed (auth) - already live")
body(doc, "Path prefixes already match each service's route namespace, so all NEXT_PUBLIC_*_API_URL "
          "point to ONE origin https://api.<domain> and Caddy routes by prefix. No per-service "
          "subdomains needed.")
body(doc, "Region: ap-south-1 (Mumbai) - Indian datasets/users (IMD, BDA, Karnataka land records).")

# 8. AWS free tier
heading(doc, "7. AWS Free Tier - Everything Available")
table(doc, ["Resource", "Free allowance", "Use here"], [
    ["New-account credits", "up to $200, 6-mo plan", "Effectively covers first ~6 mo of the t3.medium"],
    ["EC2", "750 h/mo t3.micro (12-mo)", "We run t3.medium (NOT free) -> ~$30/mo on-demand"],
    ["EBS", "30 GB gp3 (12-mo)", "Root + data volume - fits free 12 mo"],
    ["Data transfer out", "100 GB/mo always-free", "Covers low-traffic API egress"],
    ["CloudFront", "1 TB/mo + always-free", "Optional (Vercel already CDNs frontend)"],
    ["Lambda", "1M req + 400k GB-s always-free", "Not used (could offload light svcs later)"],
    ["S3", "5 GB (12-mo)", "Optional: export PDFs / static assets"],
    ["SSM Parameter Store", "standard params free", "Store secrets (NOT paid Secrets Manager)"],
    ["CloudWatch", "10 metrics, 5 GB logs, 3 dashboards", "Basic monitoring"],
    ["Route 53", "none (hosted zone $0.50/mo)", "Optional if AWS-managed DNS"],
    ["Supabase", "separate free tier", "Already in use"],
])
body(doc, "Est. monthly cost: t3.medium on-demand ~$30 + EBS ~$2.4 = ~$33/mo (inside budget). "
          "First ~6 mo can be ~$0 against the $200 credits. Cut further: stop instance when idle, "
          "or drop to t3.small (~$15) with swap + container mem limits if RAM holds.", bold=False)

# 9. Containerizing
heading(doc, "8. Containerizing - YES, Keep It")
bullets(doc, [
    "Heavy native wheels (netCDF4, scipy, xarray, pvlib, shapely, pyproj) -> reproducible only in containers; bare-host venvs are fragile across the 2 heavy services.",
    "docker-compose.yml + 10 Dockerfiles already exist and pass FVD workflow -> zero rebuild.",
    "One-command deploy + health checks + per-service mem limits.",
    "Matches CLAUDE.md migration workflow. No reason to de-containerize for a single-box beta.",
])

# 10. Deployment plan
heading(doc, "9. Deployment Plan (Ordered)")
heading(doc, "A. Prep (local, no infra spend)", level=2)
bullets(doc, [
    "Add infra/Caddyfile - reverse proxy, auto-TLS for api.<domain>, path-prefix -> service.",
    "Verify CORS on each FastAPI service allows the Vercel origin (not *). Add/centralize CORS middleware if missing.",
    "Add per-service mem_limit (+ heavy services more) to a deploy compose override; add 2 GB host swap as safety.",
    "Confirm docker compose build succeeds for all 10 locally; curl each /health.",
    "Decide domain + DNS (Route 53 or external registrar). Reserve api.<domain> for EC2 EIP.",
], numbered=True)
heading(doc, "B. Provision EC2", level=2)
bullets(doc, [
    "Launch t3.medium (ap-south-1), Ubuntu/Amazon Linux, 30 GB gp3, Elastic IP.",
    "Security group: 80/443 from anywhere; 22 from your IP only.",
    "Install Docker + compose plugin.",
    "Put secrets in SSM Parameter Store; fetch at boot into .env (gitignored). gee-sa.json mounted read-only. Copy IMD .grd to data volume; confirm it fits 30 GB (bump EBS if not).",
], numbered=True)
heading(doc, "C. Deploy backend", level=2)
bullets(doc, [
    "Get code onto box (read-only deploy key or rsync). Local-only per repo rules - no GitHub Actions.",
    "docker compose -f docker-compose.yml -f compose.deploy.yml up -d --build.",
    "Caddy fronts all services + issues TLS. Smoke each endpoint via https://api.<domain>/... -> expect 200/healthy.",
], numbered=True)
heading(doc, "D. Wire frontend (Vercel)", level=2)
bullets(doc, [
    "Import apps/web. Set env: all NEXT_PUBLIC_*_API_URL = https://api.<domain>, Supabase URL/anon key, MapTiler key.",
    "Deploy. E2E: load a project, run each analysis module against live backend; verify 2D/3D map, shadows, amenities, charts.",
], numbered=True)
heading(doc, "E. Gates + go-live", level=2)
body(doc, "Run the gate checklist (Section 10). Enable only validated FLAGS in prod .env. Tag release.")

# 11. Gates
heading(doc, "10. Deployment Gates (Local - No GitHub Actions)")
body(doc, "Pre-deploy checklist / script under scripts/:")
table(doc, ["#", "Gate", "Check"], [
    ["1", "Contract sync", "contracts/<svc>.yaml matches FastAPI routes (contract-validator agent). PRs touching contracts/ updated CHANGELOG."],
    ["2", "Flags default-off", "New behavior gated; only validated flags enabled in prod .env."],
    ["3", "Lint/format", "ruff check + ruff format on services/ + flags."],
    ["4", "Type check", "npx tsc --noEmit (web) clean; pyright optional."],
    ["5", "Tests", "pytest tests/ smoke + per-service; npm run build green."],
    ["6", "Security review", "/security-review on branch; no secrets in tracked files."],
    ["7", "Image scan", "docker build all; Trivy scan for high/critical CVEs."],
    ["8", "Health green", "every service /health passes in compose."],
    ["9", "Deployed smoke", "hit each endpoint through Caddy, assert 200."],
    ["10", "Manual approval", "flip prod FLAGS only after manual validation."],
])

# 12. API inventory
heading(doc, "11. Full API Inventory")
heading(doc, "A. External APIs the frontend calls directly", level=2)
table(doc, ["API", "Host", "Purpose", "File"], [
    ["Overpass (primary)", "overpass.openstreetmap.fr/api/interpreter", "building footprints", "lib/osm.ts:13"],
    ["Overpass (fallback)", "overpass-api.de/api/interpreter", "same, on .fr fail", "lib/osm.ts:14"],
    ["Nominatim", "nominatim.openstreetmap.org", "reverse geocode (place name)", "lib/osm.ts:15"],
    ["CARTO Voyager tiles", "basemaps.cartocdn.com", "2D raster basemap (no key)", "MapContainer.tsx:55"],
    ["MapTiler", "vector tiles (env key)", "3D basemap + building vector source", "Scene3D.tsx"],
    ["Supabase", "env URL", "auth (OAuth/PKCE)", "lib/supabase/client.ts"],
])
heading(doc, "B. Backend services (frontend -> NEXT_PUBLIC_*_API_URL)", level=2)
table(doc, ["Service", "Port", "Endpoints", "Method"], [
    ["temperature", "8000", "/weather/thermal-profile, /weather/thermal-grid", "GET, POST"],
    ["sunpath", "8001", "/sunpath/annual, /sunpath/solar-day", "GET"],
    ["flood", "8002", "/flood/analyze", "POST"],
    ["wind", "8003", "/wind/analyze", "POST"],
    ["rainfall", "8004", "/rainfall/archive, /rainfall/summary", "GET, POST"],
    ["geo", "8005", "/geo/amenities, /geo/zone, /geo/soil, /geo/water-constraints", "GET"],
    ["planning", "8006", "/planning/analyze", "POST"],
    ["infrastructure", "8007", "/infrastructure/analyze", "POST"],
    ["future-infra", "8008", "/future-infra/pipeline", "GET"],
    ["land-records", "8009", "/land-records/lookup", "POST"],
])
body(doc, "Plus /health on each (compose healthchecks).")
heading(doc, "C. Upstream APIs the backend calls", level=2)
table(doc, ["Upstream", "Used by"], [
    ["Overpass (OSM)", "geo, infrastructure, water-constraints"],
    ["Open-Meteo", "temperature, wind, rainfall"],
    ["Google Earth Engine (svc-account)", "flood, sunpath, rainfall, geo NDVI/LULC"],
    ["SoilGrids (ISRIC)", "geo /soil"],
    ["ISRO Bhuvan / NRSC LULC", "geo /zone"],
    ["CHIRPS (via GEE)", "rainfall"],
    ["IMD gridded .grd (local files, not API)", "temperature"],
])
body(doc, "Deploy impact: CORS allowlist = Vercel origin on all 10 services. Caddy routes by the "
          "unique path prefixes -> one origin works. Egress firewall must allow "
          "Overpass/Nominatim/Open-Meteo/GEE/SoilGrids/MapTiler/CARTO/Supabase.")

# 13. Env / secrets
heading(doc, "12. Environment Variables / Secrets Matrix")
table(doc, ["Var", "Where", "Holds"], [
    ["NEXT_PUBLIC_*_API_URL (x10)", "Vercel", "https://api.<domain> (one origin)"],
    ["NEXT_PUBLIC_SUPABASE_URL / ..._ANON_KEY", "Vercel", "Supabase auth"],
    ["NEXT_PUBLIC_MAPTILER_KEY", "Vercel", "3D vector tiles"],
    ["FLAGS", "EC2 .env per service", "enable validated features only"],
    ["gee-sa.json", "EC2 (SSM -> file, RO mount)", "GEE service account"],
    ["Supabase service keys (if needed)", "EC2 SSM Param Store", "server-side"],
])
body(doc, "IMD .grd data -> EBS volume mount into temperature container.")

# 14. CORS + Caddy
heading(doc, "13. CORS + Reverse-Proxy (Caddy) Notes")
bullets(doc, [
    "Each FastAPI service: CORSMiddleware allow_origins = exact Vercel domain(s), not *.",
    "infra/Caddyfile: one site api.<domain>, auto-TLS (Let's Encrypt), reverse_proxy by path prefix to each service:port.",
    "Health route / or /health for uptime checks.",
])

# 15. Monitoring / rollback
heading(doc, "14. Monitoring, Rollback, Post-Deploy")
bullets(doc, [
    "Monitoring: CloudWatch agent (CPU/RAM/disk), container docker stats, Caddy access logs.",
    "Rollback: tag images per release; docker compose down && up -d previous tag. Keep last known-good .env. Flags off = instant feature kill switch without redeploy.",
    "Post-deploy checklist: all /health green; each endpoint 200 via Caddy; frontend E2E all modules; 2D+3D map render; shadows; amenities; export; auth login.",
])

# 16. Risks
heading(doc, "15. Risks / Unknowns")
bullets(doc, [
    "RAM: t3.medium 4GB with temperature+sunpath+8 light is workable but watch scipy/xarray import spikes -> enforce mem_limit + 2 GB swap. t3.small only if monitoring confirms headroom.",
    "IMD data size - confirm .grd fits 30 GB EBS; bump if large.",
    "Single EC2 = single point of failure - acceptable for beta, no HA/autoscale.",
    "External API limits - Overpass/Open-Meteo/GEE rate limits; add response caching later.",
    "MapTiler free tier request cap - monitor 3D tile usage.",
    "CORS/HTTPS - lock CORS to Vercel domain; never * over public internet.",
    "Cost creep - egress beyond 100 GB/mo; stop instance when idle during beta.",
])

# 17. Verification
heading(doc, "16. Verification")
bullets(doc, [
    "Local: docker compose up -> curl each /health -> pytest tests/ -> npx tsc --noEmit -> npm run build.",
    "Deployed: curl https://api.<domain>/sunpath/annual?..., /weather/thermal-profile?..., etc -> 200; "
    "load Vercel app, exercise every analysis module + 2D/3D map + export.",
])

out = Path(r"C:\Users\tanny\OneDrive\Desktop\Site\SAT\docs\SAT_Deployment_Plan.docx")
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(out))
print("WROTE", out)
