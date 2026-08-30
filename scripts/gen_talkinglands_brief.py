"""Generate docs/TalkingLands_Research_Brief.docx — competitor research brief.

Team-shareable teardown of TalkingLands (spatial-intelligence proptech) to
benchmark SAT's planned Builders View. Mirrors the house style of
scripts/gen_deployment_doc.py. Source content: the approved research spike at
.claude/plans/i-need-to-understand-velvety-pearl.md (Part 1).
Run: python scripts/gen_talkinglands_brief.py
"""

from __future__ import annotations

import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# ── palette (SAT brand) ──────────────────────────────────────────────────────
BRAND = RGBColor(0x2E, 0x7D, 0x6F)      # primary teal
INK = RGBColor(0x1C, 0x24, 0x20)        # near-black
MUTED = RGBColor(0x5A, 0x6A, 0x66)
MONO_BG = "EAF2F1"                       # secondary tint


def shade(cell, hex_fill: str) -> None:
    tc = cell._tc.get_or_add_tcPr()
    shd = tc.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tc.append(shd)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = BRAND
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pbdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pbdr.makeelement(qn("w:bottom"), {})
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "2E7D6F")
        pbdr.append(bottom)
        pPr.append(pbdr)
    else:
        run.font.size = Pt(12.5)
        run.font.color.rgb = INK
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)


def body(doc: Document, text: str, bold=False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(4)


def bullets(doc: Document, items, numbered=False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for it in items:
        p = doc.add_paragraph(style=style)
        r = p.add_run(it)
        r.font.size = Pt(10.5)
        r.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(1)


def quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(10)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED
    shade_paragraph(p, MONO_BG)


def shade_paragraph(p, hex_fill: str) -> None:
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


def table(doc: Document, headers, rows) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], "2E7D6F")
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if ri % 2 == 1:
                shade(cells[i], "F4F8F7")
            para = cells[i].paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── document ─────────────────────────────────────────────────────────────────
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

# title page
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = tp.add_run("TalkingLands — Competitor Research Brief")
tr.bold = True
tr.font.size = Pt(24)
tr.font.color.rgb = BRAND
st = doc.add_paragraph()
sr = st.add_run("Spatial-intelligence proptech teardown · benchmark for SAT Builders View")
sr.font.size = Pt(12)
sr.font.color.rgb = MUTED

meta = doc.add_paragraph()
meta.paragraph_format.space_before = Pt(8)
for label, val in [
    ("Date", datetime.date.today().isoformat()),
    ("Prepared by", "Tanmay C J"),
    ("Purpose", "Understand how TalkingLands works (tech, data, sources, parcel pinpointing)"),
    ("Status", "Research spike — no product code written"),
    ("Audience", "SAT team"),
]:
    rr = meta.add_run(f"{label}:  ")
    rr.bold = True
    rr.font.size = Pt(10)
    rr.font.color.rgb = INK
    vv = meta.add_run(f"{val}\n")
    vv.font.size = Pt(10)
    vv.font.color.rgb = MUTED

# 1. Executive summary
heading(doc, "1. Executive Summary")
body(doc, "TalkingLands is a Bengaluru proptech (founded 2021, largely bootstrapped) that has built "
          "India's first parcel-level “spatial intelligence” platform for real estate. Its "
          "differentiator is not a novel map engine — it is the aggregation, georeferencing, and "
          "cleaning of official Indian land-records data (survey numbers, cadastral parcels, CDP/zoning, "
          "risk buffers) into a single map-first product.")
body(doc, "Users type a survey number and get the exact land parcel pinned on the map, wrapped in six "
          "categories of contextual layers and an AI-generated property report. Coverage is deepest in "
          "Karnataka and expanding to Telangana and Maharashtra. Monetisation is a pay-per-layer "
          "“TL Coins” credit model. This is the capability SAT lacks for a Builders audience.")

# 2. Company snapshot
heading(doc, "2. Company Snapshot")
table(doc, ["Field", "Detail"], [
    ["Founded", "2021, Bengaluru, Karnataka"],
    ["Founders", "Sriram Chitlur (CEO), Srinath Setty"],
    ["Funding", "Effectively bootstrapped — Crunchbase shows a single ~$10.9K angel round (Mar 2024); "
                "Inc42 lists it as bootstrapped"],
    ["Scale claims", "50,000+ users; ₹5,000 Cr in land decisions influenced; ₹900 Cr properties sold"],
    ["Coverage", "Deepest: Karnataka (BBMP wards, Greater Bengaluru, CDP). Expanding: Telangana, "
                 "Maharashtra. “Searchable across India where layer data exists”"],
    ["Govt traction", "Karnataka IT/BT Minister Priyank Kharge engaged them on land-governance use cases"],
    ["Partnership", "With Datasamudra (Teleindia Datacenter) + Vishanti Systems → “Bhu Drishti” "
                    "geospatial platform, Bengaluru Tech Summit 2025"],
])

# 3. Product suite
heading(doc, "3. Product Suite")
bullets(doc, [
    "Discover — consumer map-based property search & booking.",
    "Insights — per-location intelligence: survey number, GPS, land use, 6 layer categories, AI "
    "property report. Free “Property” and “Risk” consults; ₹2,500 “Rajakaluve” "
    "(storm-drain / govt-boundary) verification.",
    "Realm — end-to-end land-acquisition platform for developers: 150+ spatial layers + deal-pipeline "
    "management + AI reports in one map-first interface.",
    "Reos — real-estate sales / project-management OS (live bookings, ledgers, transactions).",
    "Real Space — drone-based virtual property tours.",
])

# 4. The map
heading(doc, "4. The Map — What It Shows and How It Works")
body(doc, "The product is a basemap + cadastral parcel overlay + thematic layers. Every revenue-village "
          "land parcel is drawn as a polygon labelled by survey number (in our captured screenshot, "
          "labels like 36/*/* and 37/*/* are survey number / hissa sub-division wildcards). Six "
          "intelligence layers sit on top, deepest in Karnataka:")
table(doc, ["Intelligence layer", "Named sub-layers"], [
    ["Growth", "National Highways, Metro & Railway, Smart City Zones, Industrial Corridors, KWIN / Bharatmala"],
    ["Boundary", "Survey Boundaries, BBMP Wards, Greater Bengaluru CDP Zones, Panchayat Limits"],
    ["Connectivity", "Airports, Metro, Railway, Bus Depots, Highways"],
    ["Amenity", "Schools, Hospitals, Shopping, Restaurants, Parks"],
    ["Risk", "Flood Index, Rajakaluves, CRZ Zones, Power Lines, Protected Areas"],
    ["Utility", "Water Pipeline, Storm Water Drain, Sewage Pipeline, Canal"],
])

# 5. Pinpoint mechanism
heading(doc, "5. How “Pinpoint per Parcel” Is Achieved")
body(doc, "A survey number maps to a georeferenced revenue-village cadastral map (the FMB / Tippan survey "
          "sketch), which renders as the parcel polygon. The national / state backbone they build on:")
bullets(doc, [
    "Bhu-Naksha — NIC's GIS cadastral engine, run by many states; renders parcel polygons by survey number.",
    "ULPIN — new 14-digit lat/long parcel ID under DILRMP (OGC-compliant), the emerging single source of truth.",
    "Dishaank (Karnataka, KSRSAC) — links historic cadastral maps to live GPS; walk onto a plot and it "
    "drops a pin with the official survey number.",
])
body(doc, "TalkingLands' real work is to ingest these per-state sources and stitch / clean them into one "
          "unified parcel database — that data aggregation is the moat, not the rendering.")

# 6. Data sources
heading(doc, "6. Data Sources")
body(doc, "In their own words:")
quote(doc, "“Government records — state revenue departments for land records, Survey of India for "
           "boundary and topographical data, municipal planning authorities for CDP and zoning maps, and "
           "government project databases.”")
body(doc, "Per-state record systems they build on (illustrative):")
table(doc, ["State / level", "Record system(s)"], [
    ["Karnataka", "Bhoomi / Dishaank / KGIS (KSRSAC)"],
    ["Maharashtra", "Mahabhulekh (7/12)"],
    ["Telangana", "Dharani"],
    ["Andhra Pradesh", "MeeBhoomi / APSAC"],
    ["Tamil Nadu", "Patta-Chitta / FMB"],
    ["Uttar Pradesh", "Bhulekh"],
    ["National", "Bhu-Naksha; Bhuvan (ISRO / NRSC) WMS"],
])

# 7. Monetisation
heading(doc, "7. Monetisation — “TL Coins” Credit Model")
body(doc, "Free to explore; 5,000 TL Coins granted on signup. Coin packs ₹100–₹1,000 "
          "(up to 50% bonus). Layers are unlocked per-view by tier:")
table(doc, ["Tier / item", "Price", "Example"], [
    ["Bronze layer", "30 coins", "Village maps"],
    ["Silver layer", "100 coins", "Greater Bengaluru zones"],
    ["Gold layer", "300 coins", "Rajakaluve risk data"],
    ["Property report", "₹100", "Full AI location report"],
    ["Rajakaluve consult", "₹2,500", "Storm-drain / govt-boundary verification"],
])
body(doc, "Note: the locked “Vinayakanagar – 30” pin in our screenshot is a 30-coin Bronze "
          "village-map layer — i.e. pay-to-unlock.")

# 8. Technology read
heading(doc, "8. Technology Read (Inferred — they publish no stack)")
bullets(doc, [
    "Standard web-GIS pattern: vector / tile basemap + parcel polygons as a vector / WMS overlay + "
    "toggleable thematic layers; survey-number → polygon lookup against a parcel DB.",
    "AI = LLM-generated natural-language property report summarising the layer data "
    "(“150+ layers decoded into one report in under a minute”).",
    "No proprietary satellite / sensor; the value is data aggregation + cleaning + UX, not raw technology.",
])

# 9. Competitive landscape
heading(doc, "9. Competitive Landscape")
table(doc, ["Company", "Focus", "Notes"], [
    ["TalkingLands", "Parcel-level spatial / map intelligence", "Bengaluru, 2021, bootstrapped, Karnataka-deep"],
    ["Landeed", "Document / title retrieval (EC, RTC, 7/12)",
     "Hyderabad, 2022, YC-backed, $16.3M, 24 states — map-light, doc-heavy (the inverse of TalkingLands)"],
    ["Teal, Zapkey, mypatta", "Title / transaction adjacents", "Smaller / niche"],
])
body(doc, "Read: the market splits into map-first (TalkingLands) vs document-first (Landeed). SAT's opening "
          "is map-first plus analysis depth (climate, zoning, 3D) that neither competitor offers.")

# 10. Caveats
heading(doc, "10. Caveats on the Claims (for the team)")
bullets(doc, [
    "Layer count: marketing says 150+; the Tracxn profile says 40+. Treat 150+ as aspirational.",
    "Funding: essentially bootstrapped — their moat is data-hustle and state relationships, not capital. "
    "SAT can compete on capability without out-spending them.",
    "Coverage outside Karnataka / Telangana / Maharashtra is thin — national parity is unproven.",
])

# 11. What it means for SAT
heading(doc, "11. What This Means for SAT")
body(doc, "SAT is strong for Architects (Rainfall, Sunpath, Wind, Temperature, Floodplains + 3D) but has no "
          "survey-number site search and no cadastral parcels — the core Builders need. The good news: "
          "the plumbing largely already exists in the codebase.")
bullets(doc, [
    "KGIS is already integrated (geo service): reverse coordinate → survey number, with land-records "
    "deep links (Bhoomi / KAVERI).",
    "A zoning / FAR / setback engine and a 3D map already exist and are reusable.",
    "The unlock: KGIS already exposes a forward service that returns a parcel polygon for a survey number "
    "(Karnataka) — the exact “type survey number → parcel on map” feature, available now.",
    "Licensing decision: spike on KGIS now; obtain a formal data license before funding; keep an "
    "own-data (digitise / aggregate) fallback ready, since KGIS terms are non-commercial.",
    "Rural-accuracy fix: add a satellite basemap option; survey-number search gives pinpoint rural "
    "locations that the current OSM/Nominatim search cannot.",
])
body(doc, "Full internal roadmap (phased build A–D, files to touch, risks) is in the research spike: "
          ".claude/plans/i-need-to-understand-velvety-pearl.md (Part 2).")

# 12. Sources
heading(doc, "12. Sources")
bullets(doc, [
    "TalkingLands site — talkinglands.com, /insights, /realm",
    "TalkingLands blog — “How to Find Survey Number of Land (2026)”",
    "Crunchbase, Tracxn, Inc42 — company / funding / profile",
    "CXO DigitalPulse — Karnataka minister engagement; The Tribune — Bhu Drishti partnership",
    "KGIS Web API (kgis.ksrsac.in); Bhu-Naksha (bhunaksha.nic.in); Bhuvan WMS (bhuvan-vec1.nrsc.gov.in)",
    "DILRMP / ULPIN (dilrmp.gov.in, dolr.gov.in/ulpin)",
    "Landeed (landeed.com, Crunchbase, Y Combinator)",
])

out = Path(r"C:\Users\tanny\OneDrive\Desktop\Site\SAT\docs\TalkingLands_Research_Brief.docx")
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(out))
print("WROTE", out)
