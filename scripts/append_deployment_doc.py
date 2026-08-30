"""Append the image-delivery / Pages comparison (Sections 17-23) to the EXISTING
Downloads copy of SAT_Deployment_Plan.docx. Does not regenerate from scratch, so
nothing already in the file is lost.

Registry decision: use GHCR (GitHub Container Registry) for image delivery, NOT
Docker Hub. Section 17b lists why Docker Hub is a poor fit and gives the GHCR method.

Run: python scripts/append_deployment_doc.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

TARGET = Path(r"C:\Users\tanny\Downloads\SAT_Deployment_Plan.docx")

BRAND = RGBColor(0x2E, 0x7D, 0x6F)
INK = RGBColor(0x1C, 0x24, 0x20)
MUTED = RGBColor(0x5A, 0x6A, 0x66)
MONO_BG = "EAF2F1"


def shade(cell, hex_fill: str) -> None:
    tc = cell._tc.get_or_add_tcPr()
    shd = tc.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tc.append(shd)


def shade_paragraph(p, hex_fill: str) -> None:
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


def heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = BRAND
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pbdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pbdr.makeelement(qn("w:bottom"), {})
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "2E7D6F")
        pbdr.append(bottom)
        pPr.append(pbdr)
    else:
        run.font.size = Pt(12.5)
        run.font.color.rgb = INK
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)


def body(doc, text: str, bold=False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(4)


def bullets(doc, items, numbered=False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for it in items:
        p = doc.add_paragraph(style=style)
        r = p.add_run(it)
        r.font.size = Pt(10.5)
        r.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(1)


def mono(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(6)
    for line in text.split("\n"):
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = INK
    shade_paragraph(p, MONO_BG)


def table(doc, headers, rows) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], "2E7D6F")
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if ri % 2 == 1:
                shade(cells[i], "F4F8F7")
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── open existing doc ────────────────────────────────────────────────────────
if not TARGET.exists():
    raise SystemExit(f"Target not found: {TARGET}")
doc = Document(str(TARGET))

# page break + addendum banner
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
banner = doc.add_paragraph()
brun = banner.add_run("ADDENDUM — Image Delivery (Registry), Docker Hub vs GHCR & \"Pages\"")
brun.bold = True
brun.font.size = Pt(18)
brun.font.color.rgb = BRAND
sub = doc.add_paragraph()
srun = sub.add_run("Budget hard cap: $30/mo for ONE t3.medium, zero other paid services.")
srun.italic = True
srun.font.size = Pt(10.5)
srun.font.color.rgb = MUTED

# 17
heading(doc, "17. Image Delivery — Build-on-Box vs a Registry")
body(doc, "The original plan (Section 9C) clones the repo to the EC2 and runs docker compose "
          "build ON the box. On a 2 vCPU / 4 GB t3.medium, building the 2 heavy images "
          "(scipy/numpy/netCDF4/shapely/pyproj/matplotlib wheels) risks OOM mid-build, takes "
          "10-20+ minutes, competes with running services, and dumps source + build cache onto "
          "EBS. Better: build OFF the box, push the built images to a registry, and pull them on "
          "the EC2. A registry is just an image store - the images still RUN on the EC2.")
body(doc, "Key fact: the SAT GitHub repo is PUBLIC, so public images expose nothing new "
          "(secrets are runtime-mounted, never baked into images). Public registries are fine.")
table(doc, ["Option", "$/mo", "Notes", "Verdict"], [
    ["Build on EC2 (current)", "0", "uses EBS + build cache", "NO - OOM risk on heavy, slow, repo+cache on box"],
    ["Docker Hub Free", "0", "1 private repo only; pull rate limits; 6-mo purge; separate account", "WEAK - see 17b"],
    ["GHCR (ghcr.io) public", "0", "uses existing GitHub; unlimited public; incremental layers", "USE THIS (17b)"],
    ["AWS ECR", ">0 after 500 MB/12mo", "500 MB free << 4 GB of images", "NO - breaks the $30-only cap"],
    ["docker save | gzip | ssh -> load", "0", "no registry/account; private; manual ~4 GB/deploy", "FALLBACK - single box only"],
    ["Self-host registry on EC2", "0", "uses EBS, extra moving part", "NO - circular"],
])

# 17b
heading(doc, "17b. Why Docker Hub Is a Poor Fit — Use GHCR Instead", level=2)
body(doc, "We mean using Docker Hub as the image REGISTRY (the store we push to / pull from), not "
          "as a host. Even for that, the free tier fits SAT badly:")
bullets(doc, [
    "Only 1 private repo on the free tier - SAT has 10 service images, so you are forced to either go fully public or pay. (We can go public, but it should be a choice, not a constraint.)",
    "Pull rate limits: anonymous ~100 pulls / 6h per IP, authenticated free ~200 / 6h. A redeploy or a box reboot pulls 10 images at once and restarts can burn through this - throttled pulls stall the deploy.",
    "6-month inactivity retention: free-tier images that are not pulled for 6 months can be purged - your release images can silently vanish.",
    "Separate account + access token to create, rotate, and store on the EC2 - one more secret unrelated to anything else you run.",
    "Namespace sprawl: 10 standalone Docker Hub repos to name and manage, disconnected from the source code and releases.",
])
body(doc, "Use GHCR (GitHub Container Registry, ghcr.io) instead - same push/pull workflow, "
          "none of the above problems:", bold=True)
bullets(doc, [
    "Uses the GitHub account that ALREADY hosts SAT - no new account, no extra service.",
    "Free and unlimited for public images (and 500 MB free if you later make some private).",
    "Generous limits, no Docker-Hub-style pull throttling or 6-month purge of release tags.",
    "Incremental layer transfer: only changed layers move on each deploy -> fast iterative beta pushes.",
    "One GitHub PAT (write:packages) you already manage; push LOCALLY -> no GitHub Actions (honours the repo's local-only rule).",
    "Images live next to the code + releases (provenance), e.g. ghcr.io/<owner>/sat-<svc>:<gitsha>.",
])
body(doc, "GHCR method (locked):", bold=True)
bullets(doc, [
    "Local login: echo $PAT | docker login ghcr.io -u <gh-user> --password-stdin.",
    "compose.deploy.yml uses image: ghcr.io/<owner>/sat-<svc>:<gitsha> (not build:); docker compose build && docker compose push.",
    "On EC2: docker login ghcr.io (read-only PAT) then docker compose -f docker-compose.yml -f compose.deploy.yml pull && up -d.",
    "Rollback: pull the previous :<gitsha> tag and up -d.",
    "No-registry fallback (single box, fully private, $0): docker save <imgs> | gzip | ssh ec2 'docker load'.",
])

# 18
heading(doc, "18. \"Deploy as Pages\" — What Is Actually Possible")
body(doc, "A registry (Docker Hub or GHCR) is not a host: it stores images, it does not run them. "
          "So neither frontend nor backend can be \"deployed on\" a registry - they run on the EC2 "
          "(backend) and Vercel (frontend). Separately, real \"Pages\" hosts (GitHub Pages, "
          "Cloudflare Pages, Netlify, Vercel) serve static HTML/CSS/JS free, with no server runtime:")
body(doc, "Frontend as a page:", bold=True)
bullets(doc, [
    "Next.js 16 app CANNOT pure-static-export. Blockers: /project/[id] and /project/[id]/export are runtime user-created dynamic routes with no generateStaticParams and no output:'export' (next.config.ts:3, app/project/[id]/page.tsx:1).",
    "So GitHub Pages = not viable without route rework. Auth + maps are already client-side (ssr:false) so those do not block.",
    "Hosts that DO run Next SSR free: Vercel (zero rework, LOCKED) or Cloudflare Pages (next-on-pages adapter, edge-runtime constraints).",
])
body(doc, "Backend as a page:", bold=True)
bullets(doc, [
    "Impossible - Pages serve static assets, not a running Python server.",
    "Every free serverless/PaaS fails the 2 heavy services: Cloudflare Workers (Python beta/Pyodide - no netCDF4/scipy/GEE native wheels), Vercel Python Functions (250 MB limit, cold starts), Render free (512 MB RAM, spins down, 1 svc), Fly.io / Koyeb (tiny free, 1-2 machines).",
    "10 services x native wheels + IMD data + GEE -> no free home. Heavy services force the EC2.",
])
body(doc, "Both free? No. Frontend free is real; the heavy backend cannot be free -> the $30 EC2 stays.", bold=True)

# 19
heading(doc, "19. Free-Tier Storage Breakdown (Extreme Detail)")
table(doc, ["Item", "Size", "Lands on"], [
    ["temperature image", "~1.2 GB", "EBS once pulled/built"],
    ["sunpath image", "~1.4 GB", "EBS"],
    ["8 light images (~250 MB each)", "~2.0 GB", "EBS"],
    ["Caddy image", "~50 MB", "EBS"],
    ["TOTAL images", "~4.0-4.6 GB", "EBS"],
    ["IMD .grd data (NOT in repo, runtime mount)", "UNKNOWN - must measure", "EBS data dir"],
    ["geo GeoJSON (baked in image)", "266 KB", "image"],
    ["future-infra config (baked in image)", "8 KB", "image"],
    ["OS + Docker engine", "~3 GB", "EBS root"],
    ["Container writable layers + logs", "~1-2 GB", "EBS"],
])
body(doc, "EBS = 30 GB gp3 (free 12 mo). Fits comfortably IF IMD data < ~18 GB (confirm by "
          "measuring services/temperature/data before deploy). Registry choice does NOT change "
          "on-box EBS - images land on EBS once pulled regardless. GHCR storage is off-box and "
          "free at this size, so the registry saves build pain + gives rollback, not dollars "
          "(EC2 cost is identical either way).")

# 20
heading(doc, "20. Comparison vs the EC2-Build Approach")
table(doc, ["Dimension", "Current (build-on-box)", "Final (build-off-box + GHCR pull)"], [
    ["Cost", "$30 (EC2)", "$30 (EC2) - identical"],
    ["Heavy-image build", "on 4 GB box -> OOM risk, slow", "on dev machine -> safe, fast"],
    ["Deploy speed", "rebuild each time", "pull changed layers (seconds)"],
    ["Rollback", "rebuild old commit", "re-tag pull (seconds)"],
    ["EBS pressure", "images + source + build cache", "images only (no source/cache)"],
    ["Registry friction", "n/a", "GHCR: existing GitHub acct, no rate-limit/purge"],
    ["Frontend", "Vercel free", "Vercel Hobby (LOCKED)"],
])
body(doc, "Net: same $30, more reliable + faster + cleaner rollback. \"Both free\" is not "
          "achievable for the heavy backend. Adopt build-off-box delivery via GHCR; keep the "
          "single $30 EC2 for all 10 services.")
body(doc, "Frontend host nuance: Vercel Hobby is non-commercial per ToS. LOCKED decision = "
          "A. Vercel Hobby (accept ToS for the private beta). Alternatives if commercial: "
          "B. Next.js on the same $30 EC2 as a container behind Caddy ($0 extra, +~400 MB RAM, "
          "you manage it); C. Cloudflare Pages (free, commercial-OK, adapter complexity).")

# 21
heading(doc, "21. Simple Explanation (Analogy)")
bullets(doc, [
    "Doc's current approach (build-on-box): ship raw ingredients (source code) to a small restaurant kitchen (the 4 GB EC2) and cook everything there. The 2 heavy dishes (temperature, sunpath) can overwhelm the small kitchen -> it stalls / catches fire (OOM), and cooking is slow while customers wait.",
    "Registry approach (build-off-box): cook the dishes at home (your powerful PC), put them in a shared fridge (an image registry - we use GHCR). The restaurant just reheats (pull + run). Fast, safe, and if today's dish is bad you instantly serve yesterday's (rollback by image tag).",
    "The registry is the fridge, not the restaurant - it stores meals, never serves customers. The EC2 is still the restaurant that serves (runs) them.",
    "Why GHCR over Docker Hub as the fridge: it's the fridge you already own (same GitHub), with no usage meter on the door (no pull-rate-limit) and it won't throw out food left for 6 months (no purge).",
])

# 22
heading(doc, "22. Final Recommended Deployment Plan (LOCKED)")
body(doc, "Cost = $30/mo, nothing else. Frontend free on Vercel Hobby. Backend on one t3.medium.", bold=True)
bullets(doc, [
    "Frontend: Vercel Hobby. Env: all NEXT_PUBLIC_*_API_URL = https://api.<domain>, Supabase URL/anon key, MapTiler key. Zero rework (Next SSR native).",
    "Backend: 1x EC2 t3.medium (ap-south-1), 30 GB gp3 EBS (free 12 mo), Elastic IP. Runs all 10 services + Caddy via docker-compose. No building on the box.",
])
body(doc, "Image delivery (GHCR public):", bold=True)
bullets(doc, [
    "Local: docker login ghcr.io; docker compose build; tag ghcr.io/<owner>/sat-<svc>:<gitsha>; docker compose push.",
    "EC2: docker login ghcr.io (read-only PAT); docker compose -f docker-compose.yml -f compose.deploy.yml pull && up -d.",
    "compose.deploy.yml override = image: tags (not build:), remove dev bind-mounts, add per-service mem_limit + restart: unless-stopped. Heavy services get more RAM headroom.",
], numbered=True)
bullets(doc, [
    "Secrets/data: SSM Parameter Store -> .env on box at boot; gee-sa.json read-only mount; copy IMD .grd to EBS data dir - measure size first (< ~18 GB to fit 30 GB EBS).",
    "TLS + routing: Caddy auto-TLS for api.<domain>, path-prefix -> each service:port.",
    "Gates (local, no Actions): contract-sync, ruff, tsc --noEmit, pytest, /security-review, Trivy scan, all /health green, deployed smoke (each endpoint 200 via Caddy), flip prod FLAGS only after manual validation - run BEFORE docker push.",
    "Rollback: pull previous image tag + up -d; flags-off = instant feature kill switch.",
    "Storage: images ~4-4.6 GB + IMD data (TBD) + OS ~3 GB + layers ~1-2 GB on 30 GB EBS (free).",
])

# 23
heading(doc, "23. Why Docker Hub Falls Short as Our Registry — and What We Use Instead")
body(doc, "To be clear: this is about USING Docker Hub as the registry (push/pull image store), "
          "not hosting the app on it. As the registry for SAT, the Docker Hub free tier is a poor "
          "fit for concrete reasons:")
bullets(doc, [
    "1 private repo only - 10 service images force public-or-pay (should be a choice, not forced).",
    "Pull rate limits (~100 anon / ~200 authed per 6h) - a 10-image redeploy or reboot can be throttled mid-deploy.",
    "6-month inactivity purge - release images that aren't pulled can be deleted.",
    "Separate account + token to manage and store on the EC2.",
    "10 disconnected repos, not linked to the GitHub source/releases.",
])
body(doc, "Use GHCR (GitHub Container Registry) instead - same docker push/pull, but it reuses the "
          "GitHub account that already hosts SAT, is free + unlimited for public images, has no "
          "Docker-Hub pull-throttle or 6-month purge, transfers only changed layers, and pushes "
          "locally with one PAT (no GitHub Actions). For a single box, docker save + scp is a valid "
          "zero-registry, fully-private fallback.", bold=True)
oneliner = doc.add_paragraph()
orun = oneliner.add_run("In one line: don't fight Docker Hub's free-tier limits - push to GHCR "
                        "(the registry tied to the repo you already have) and run the images on the $30 EC2. "
                        "Vercel = the free storefront (frontend).")
orun.bold = True
orun.font.size = Pt(10.5)
orun.font.color.rgb = BRAND

try:
    doc.save(str(TARGET))
    print("APPENDED ->", TARGET)
except PermissionError:
    fallback = TARGET.with_name("SAT_Deployment_Plan_final.docx")
    doc.save(str(fallback))
    print("LOCKED (Word open). Saved to ->", fallback)
