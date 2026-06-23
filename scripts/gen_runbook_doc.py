# Generates docs/SAT_Deployment_Runbook.docx — Vercel frontend + EC2 t3.medium runbook.
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL = RGBColor(0x2E, 0x7D, 0x6F)
TEAL_HEX = "2E7D6F"
TINT_HEX = "EAF2F1"
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)
MONO_BG = "F4F4F4"

OUT = Path(r"C:\Users\tanny\OneDrive\Desktop\Site\SAT\docs\SAT_Deployment_Runbook.docx")

doc = Document()
styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10.5)


def shade(cell, hex_color):
    tc = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_color)
    tc.append(sh)


def heading(text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = TEAL if level == 1 else DARK
    r.font.size = Pt(16 if level == 1 else 12.5 if level == 2 else 11)
    p.space_before = Pt(10)
    p.space_after = Pt(4)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pb = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), TEAL_HEX)
        pb.append(bottom)
        pPr.append(pb)
    return p


def body(text, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(10.5)
    if color:
        r.font.color.rgb = color
    p.space_after = Pt(4)
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)
        p.paragraph_format.space_after = Pt(2)


def mono(code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1)
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts")) or OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")
    rpr.append(rfonts)
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), MONO_BG)
    pPr.append(sh)
    p.space_after = Pt(6)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], TEAL_HEX)
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if ri % 2 == 1:
                shade(cells[i], TINT_HEX)
            para = cells[i].paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().space_after = Pt(2)
    return t


# ---------- TITLE ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("SAT Deployment — Implementation Runbook")
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = TEAL
sub = doc.add_paragraph()
r = sub.add_run("Approved scope: Frontend on Vercel  +  Backend on AWS EC2 t3.medium")
r.font.size = Pt(12); r.font.color.rgb = GREY
meta = doc.add_paragraph()
r = meta.add_run("Site Analysis Tool  ·  Beta  ·  Branch ux/phase-1-research  ·  2026-06-20")
r.font.size = Pt(9.5); r.font.color.rgb = GREY

# ---------- CONTEXT ----------
heading("1. Context", 1)
body("Beta-ready stack (Phase 10 GO). Two deployment items approved to stand up now:")
bullets([
    ("Frontend → Vercel. ", "Next.js 16, free Hobby tier."),
    ("Backend → one AWS EC2 t3.medium. ", "~$30/mo, ap-south-1 (Mumbai), runs all 10 FastAPI services + Caddy reverse proxy via docker-compose."),
])
body("Image-registry choice (GHCR / Docker Hub) is deferred / not approved — baseline here builds images ON the box with swap. API domain decision: OWN DOMAIN (Caddy auto-TLS via Let's Encrypt).")

# ---------- CORRECTIONS ----------
heading("2. Corrections Found in Code", 1)
body("Verified against the live frontend + compose. The earlier draft plan was wrong on these:")
table(
    ["Item", "Correct value / behaviour", "Source"],
    [
        ["Supabase key var", "NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY (NOT ..._ANON_KEY)", "lib/env.ts:19, lib/supabase/client.ts:9"],
        ["Frontend API env vars", "10 distinct per-service vars; in prod all point at the SAME origin (Caddy routes by path)", "lib/api/analysis.ts:37-49"],
        ["Request paths", "Each call already carries its service prefix → single-origin path routing valid", "analysis.ts:152,335,601,730,832,1201,1357,1398,1417"],
        ["temperature prefix", "/weather  (NOT /temperature)", "analysis.ts:336,456"],
        ["CORS", "Env-driven per service via CORS_ORIGINS; today inconsistent (:5173 vs :3000) — override in prod", "docker-compose.yml"],
        ["Flag gating", "Each service 403s unless its FLAGS value is set", "analysis.ts:6-11"],
        ["Dev compose", "Bind mounts + --reload on every service → needs standalone prod compose (overrides can't remove mounts)", "docker-compose.yml"],
        ["Static export", "No output:'export' in next.config.ts → SSR confirmed, native Vercel fit", "next.config.ts"],
    ],
    widths=[1.5, 3.6, 1.9],
)

# ---------- PREREQS ----------
heading("3. Prerequisites", 1)
bullets([
    ("GitHub. ", "Repo Site-Analysis/SAT is public → EC2 can git clone over HTTPS. Vercel imports via GitHub connection."),
    ("Secrets/data NOT on this machine ", "(confirmed absent): gee-sa.json, root .env, IMD .grd data. Source from the Site Analysis workspace before Part B."),
    ("Measure IMD .grd size ", "— must fit 30 GB EBS (target < ~18 GB)."),
    ("AWS account ", "(ap-south-1). New account = up to $200 credits cover first ~6 months."),
    ("Domain in hand ", "(registrar access to add a DNS A record)."),
])

# ---------- PART A ----------
heading("PART A — Frontend on Vercel", 1)

heading("A1. Import project", 2)
bullets([
    "vercel.com → Add New → Project → import Site-Analysis/SAT from GitHub.",
    "Root Directory = apps/web (monorepo). Framework auto-detects Next.js.",
    "Build = next build (default). Install = npm install. Leave Output default.",
])

heading("A2. Set environment variables", 2)
body("Project → Settings → Environment Variables (scope = Production + Preview). API vars all point at the one Caddy origin https://api.<domain>:")
table(
    ["Variable", "Value"],
    [
        ["NEXT_PUBLIC_SUPABASE_URL", "https://<ref>.supabase.co"],
        ["NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY", "sb_publishable_<key>"],
        ["NEXT_PUBLIC_MAPTILER_KEY", "<maptiler key>"],
        ["NEXT_PUBLIC_TEMPERATURE_API_URL", "https://api.<domain>"],
        ["NEXT_PUBLIC_SUNPATH_API_URL", "https://api.<domain>"],
        ["NEXT_PUBLIC_FLOOD_API_URL", "https://api.<domain>"],
        ["NEXT_PUBLIC_WIND_API_URL", "https://api.<domain>"],
        ["NEXT_PUBLIC_RAINFALL_API_URL", "https://api.<domain>"],
        ["NEXT_PUBLIC_GEO_API_URL", "https://api.<domain>"],
        ["NEXT_PUBLIC_PLANNING_API_URL", "https://api.<domain>"],
        ["NEXT_PUBLIC_INFRA_API_URL", "https://api.<domain>"],
        ["NEXT_PUBLIC_FUTURE_INFRA_API_URL", "https://api.<domain>"],
        ["NEXT_PUBLIC_LAND_RECORDS_API_URL", "https://api.<domain>"],
        ["NEXT_PUBLIC_API_BASE_URL", "https://api.<domain> (verify usage; blank if same-origin Next routes)"],
    ],
    widths=[3.2, 3.6],
)

heading("A3. First deploy", 2)
bullets([
    "Deploy → get https://<project>.vercel.app. RECORD this exact origin — it's the CORS allowlist value for Part B (B4/B5).",
    "Optional: add a custom frontend domain later (Vercel → Domains). Not required for beta.",
])
heading("A4. Hold E2E", 2)
body("Hold end-to-end testing until Part B is live (frontend can't reach the API yet). Return after B8.")

# ---------- PART B ----------
heading("PART B — Backend on EC2 t3.medium", 1)

heading("B1. Provision the instance (Console → EC2, region ap-south-1)", 2)
bullets([
    "AMI: Ubuntu 22.04 LTS (x86_64). Type: t3.medium (2 vCPU / 4 GB).",
    "Storage: 30 GB gp3 root.",
    "Key pair: create/download .pem (SSH).",
])
body("Security group:")
table(
    ["Port", "Source", "Purpose"],
    [
        ["22/tcp", "your IP only", "SSH admin"],
        ["80/tcp", "0.0.0.0/0", "Let's Encrypt HTTP-01 + redirect"],
        ["443/tcp", "0.0.0.0/0", "HTTPS API"],
        ["8000-8009", "DO NOT OPEN", "internal only — Caddy is sole public entry"],
    ],
    widths=[1.3, 2.2, 3.3],
)
body("Allocate an Elastic IP → associate to the instance (stable IP for DNS).")

heading("B2. Base OS + Docker + swap", 2)
mono(
    "ssh -i sat.pem ubuntu@<elastic-ip>\n"
    "sudo apt-get update && sudo apt-get -y upgrade\n"
    "# Docker engine + compose plugin\n"
    "curl -fsSL https://get.docker.com | sudo sh\n"
    "sudo usermod -aG docker ubuntu && newgrp docker\n"
    "# 4 GB swap — protects the 4 GB box during heavy image builds (free)\n"
    "sudo fallocate -l 4G /swapfile && sudo chmod 600 /swapfile\n"
    "sudo mkswap /swapfile && sudo swapon /swapfile\n"
    "echo '/swapfile none swap sw 0 0' | sudo tee -a /etc/fstab"
)

heading("B3. Get the code", 2)
mono(
    "sudo mkdir -p /opt/sat && sudo chown ubuntu:ubuntu /opt/sat\n"
    "git clone https://github.com/Site-Analysis/SAT.git /opt/sat\n"
    "cd /opt/sat"
)

heading("B4. Place secrets + data (NOT in git — copy via scp)", 2)
mono(
    "# from your local/workspace machine:\n"
    "scp -i sat.pem gee-sa.json ubuntu@<eip>:/opt/sat/gee-sa.json\n"
    "scp -i sat.pem -r imd-data/  ubuntu@<eip>:/opt/sat/imd-data/   # the .grd files"
)
body("On the box create /opt/sat/.env (gitignored) holding non-public config:")
mono(
    "FLAGS=feature.temperature.thermal-profile,feature.sunpath.diagram,\\\n"
    "      feature.flood.risk-analysis,feature.wind.analysis,feature.rainfall.summary\n"
    "      # + any geo/planning/infra/future-infra/land flags validated\n"
    "CORS_ORIGINS=https://<project>.vercel.app\n"
    "OVERPASS_URL=https://overpass-api.de/api/interpreter"
)
body("Optional (still free): store these in SSM Parameter Store and fetch at boot instead of a flat file. Skip for first beta.")

heading("B5. Author the prod compose + Caddyfile (new files)", 2)
body("infra/docker-compose.prod.yml — standalone (do not layer over dev compose):")
bullets([
    "All 10 services with build: ./services/<svc> (build-on-box baseline).",
    "Remove the ./services/<svc>:/app bind mounts and --reload; keep only gee-sa.json:/app/gee-sa.json:ro (sunpath/flood/wind) and /opt/sat/imd-data:/app/data:ro (temperature).",
    "Do NOT publish 8000-8009 to the host (internal network only).",
    "restart: unless-stopped; mem_limit per service (heavy temperature/sunpath ~1.2-1.5 GB each, light ~256-384 MB) — keep total under ~3.5 GB + swap.",
    "environment: CORS_ORIGINS=${CORS_ORIGINS} + FLAGS=${FLAGS} from /opt/sat/.env.",
    "Add a caddy service on the same network: image caddy:2, publish 80:80 + 443:443, mount ./infra/Caddyfile:/etc/caddy/Caddyfile:ro + named volumes caddy_data (certs) + caddy_config.",
])
body("infra/Caddyfile — one site, auto-TLS, path-prefix → service (names resolve on the compose network):")
mono(
    "api.<domain> {\n"
    "    encode gzip\n"
    "    reverse_proxy /weather/*         temperature:8000\n"
    "    reverse_proxy /sunpath/*         sunpath:8001\n"
    "    reverse_proxy /flood/*           flood:8002\n"
    "    reverse_proxy /wind/*            wind:8003\n"
    "    reverse_proxy /rainfall/*        rainfall:8004\n"
    "    reverse_proxy /geo/*             geo:8005\n"
    "    reverse_proxy /planning/*        planning:8006\n"
    "    reverse_proxy /infrastructure/*  infrastructure:8007\n"
    "    reverse_proxy /future-infra/*    future-infra:8008\n"
    "    reverse_proxy /land-records/*    land-records:8009\n"
    "}"
)

heading("B6. DNS (registrar)", 2)
bullets([
    "A record: api.<domain> → <elastic-ip>, TTL low (300s) for first cutover.",
    "Wait for propagation (dig api.<domain> resolves to the EIP).",
])

heading("B7. Build + bring up", 2)
mono(
    "cd /opt/sat\n"
    "# Build heavy images first (watch memory; swap covers spikes)\n"
    "docker compose -f infra/docker-compose.prod.yml --env-file .env build temperature sunpath\n"
    "docker compose -f infra/docker-compose.prod.yml --env-file .env build   # the rest\n"
    "docker compose -f infra/docker-compose.prod.yml --env-file .env up -d\n"
    "docker compose -f infra/docker-compose.prod.yml ps   # all healthy\n"
    "docker stats --no-stream                              # confirm RAM headroom"
)
body("Caddy auto-issues the TLS cert on first hit to api.<domain> (port 80 must be open — B1).")

heading("B8. Smoke (deployed)", 2)
bullets([
    "Internal (on box) per service: docker compose ... exec <svc> curl -fsS localhost:<port>/health → all 200.",
    "Through Caddy (real endpoints, prove routing + TLS):",
])
mono(
    "curl -fsS \"https://api.<domain>/sunpath/annual?lat=12.97&lon=77.59\"\n"
    "curl -fsS \"https://api.<domain>/geo/zone?lat=12.97&lon=77.59&radius_m=500\""
)
body("Expect HTTP 200 JSON (not 403 → confirms FLAGS set; not 502 → confirms upstream up).")

# ---------- CUTOVER ----------
heading("4. Cutover + Verify (ties A and B)", 1)
bullets([
    "Confirm CORS_ORIGINS in /opt/sat/.env exactly equals the Vercel origin from A3; up -d to apply if changed.",
    "Open the Vercel app → log in (Supabase) → create/open a project.",
    "Exercise EACH analysis module against live API: flood, sunpath (2D + 3D map, shadows), wind, temperature/climate, rainfall, geo/amenities, planning, infrastructure, future-infra, land-records.",
    "Verify 2D Leaflet + 3D MapLibre render, shadows, amenity markers, recharts, PDF export.",
    "Devtools Network: calls hit https://api.<domain>/<prefix>/..., status 200, no CORS errors, no mixed-content warnings.",
])

# ---------- RISKS ----------
heading("5. Risks / Watch-items", 1)
table(
    ["Risk", "Mitigation"],
    [
        ["Build-on-box OOM (2 heavy images on 4 GB)", "4 GB swap + sequential build + mem_limit. If still OOM, fall back to build-off-box (deferred GHCR) — flag to user, don't self-approve."],
        ["IMD data size", "Measure before scp; bump EBS if > ~18 GB (small added cost — surface first, breaks '$30 only')."],
        ["FLAGS coverage", "Confirm geo/planning/infra/future-infra/land flag strings; a missing flag = 403. Pull exact names from packages/flags/src/flags.py."],
        ["NEXT_PUBLIC_API_BASE_URL", "Verify what client.ts calls; if only same-origin Next routes, leave blank."],
        ["CORS default drift (:5173 vs :3000)", "Prod value from .env CORS_ORIGINS; prod compose must override hardcoded environment: on temperature/rainfall."],
        ["Single EC2 = no HA", "Acceptable for beta."],
        ["Vercel Hobby non-commercial (ToS)", "Accepted for private beta (locked decision)."],
    ],
    widths=[2.6, 4.2],
)

# ---------- FILES ----------
heading("6. Files the Implementer Will Create (Part B)", 1)
bullets([
    "infra/docker-compose.prod.yml (new)",
    "infra/Caddyfile (new)",
    "/opt/sat/.env on the box (not committed)",
    "Optional scripts/ec2-bootstrap.sh capturing B2 steps",
])

heading("7. Out of Scope (deferred — do NOT do without explicit approval)", 1)
bullets([
    "GHCR / Docker Hub image registry + build-off-box pipeline.",
    "GitHub Actions / any CI (repo rule: local-only).",
    "SSM Parameter Store wiring (optional optimization).",
    "Frontend custom domain, autoscaling, monitoring stack.",
])

OUT.parent.mkdir(parents=True, exist_ok=True)
try:
    doc.save(str(OUT))
    print("SAVED ->", OUT)
except PermissionError:
    fb = OUT.with_name("SAT_Deployment_Runbook_final.docx")
    doc.save(str(fb))
    print("LOCKED (Word open). Saved to ->", fb)

d2 = Document(str(OUT if OUT.exists() else OUT.with_name("SAT_Deployment_Runbook_final.docx")))
print("paras", len(d2.paragraphs), "| tables", len(d2.tables))
