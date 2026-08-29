"""
Generate Builder Module Technical Handoff Word Document.
Output: Desktop/Builder_Module_Handoff_Qnit.docx
"""
import subprocess, sys, os

# Auto-install python-docx if missing
try:
    from docx import Document
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document

from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Qnit palette ──────────────────────────────────────────────────────────────
GREEN       = RGBColor(0x30, 0x62, 0x23)   # #306223
GREEN_LIGHT = RGBColor(0x4A, 0x7D, 0x3B)   # lighter for table rows
GREEN_BG    = RGBColor(0xE8, 0xF0, 0xE5)   # very light green for alt rows
CREAM       = RGBColor(0xFD, 0xFC, 0xFB)   # #FDFCFB
DARK        = RGBColor(0x1A, 0x20, 0x10)   # #1a2010
MUTED       = RGBColor(0x7B, 0x8F, 0x83)   # #7B8F83
AMBER       = RGBColor(0xB4, 0x53, 0x09)   # #B45309
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BORDER_CLR  = RGBColor(0xCF, 0xD6, 0xC4)   # light sage

OUT_PATH = r"C:\Users\tanny\OneDrive\Desktop\Builder_Module_Handoff_Qnit.docx"

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def cell_para(cell, text, bold=False, color=DARK, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return run

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = GREEN
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = GREEN
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(10)
        run.font.color.rgb = MUTED
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
    # Bottom border for h1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "306223")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def add_body(doc, text, color=DARK, size=10, italic=False, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.italic = italic
    run.font.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.clear()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    # Shade background via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F4EE")
    pPr.append(shd)
    for line in text.split("\n"):
        if p.runs:
            p.add_run("\n" + line)
        else:
            run = p.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
            run.font.color.rgb = DARK
            continue
        p.runs[-1].font.name = "Courier New"
        p.runs[-1].font.size = Pt(8.5)
        p.runs[-1].font.color.rgb = DARK

def add_bullet(doc, text, color=DARK, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    # Re-apply list style via direct XML (avoids missing style errors)
    p2 = doc.add_paragraph()
    p2.clear()
    p2.paragraph_format.left_indent = Cm(0.6)
    p2.paragraph_format.space_after = Pt(3)
    run = p2.add_run("• " + text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    # Remove the empty list bullet paragraph
    p._element.getparent().remove(p._element)
    return p2

def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CFD6C4")
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Qnit  ·  Confidential  ·  qnit.site  ·  Builder Module Technical Handoff")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED

def add_page_break(doc):
    doc.add_page_break()

# ── Build Document ────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

add_footer(doc)

# ════════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════════

# Green header bar (table trick: 1 col, 1 row)
cover_bar = doc.add_table(rows=1, cols=1)
cover_bar.alignment = WD_TABLE_ALIGNMENT.CENTER
bar_cell = cover_bar.cell(0, 0)
set_cell_bg(bar_cell, GREEN)
bar_cell.width = Inches(6.3)
bar_p = bar_cell.paragraphs[0]
bar_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
bar_p.paragraph_format.space_before = Pt(18)
bar_p.paragraph_format.space_after = Pt(18)
bar_run = bar_p.add_run("QNIT")
bar_run.font.name = "Calibri"
bar_run.font.size = Pt(28)
bar_run.font.bold = True
bar_run.font.color.rgb = WHITE
bar_run2 = bar_p.add_run("  ·  Site Analysis Tool")
bar_run2.font.name = "Calibri"
bar_run2.font.size = Pt(14)
bar_run2.font.color.rgb = RGBColor(0xC8, 0xD8, 0xC2)

doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Builder Module")
title_run.font.name = "Calibri"
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = DARK

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("Technical Handoff & Data Sources")
sub_run.font.name = "Calibri"
sub_run.font.size = Pt(16)
sub_run.font.color.rgb = MUTED

doc.add_paragraph()
doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta_p.add_run("Date: 2026-08-29    ·    Contract Version: v2.32.0    ·    Confidential")
meta_run.font.name = "Calibri"
meta_run.font.size = Pt(10)
meta_run.font.color.rgb = MUTED

doc.add_paragraph()

desc_p = doc.add_paragraph()
desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_run = desc_p.add_run(
    "This document covers all data sources, API endpoints, Supabase setup,\n"
    "feature flags, and developer handoff notes for the Builder Module\n"
    "of the Qnit Site Analysis Tool (SAT)."
)
desc_run.font.name = "Calibri"
desc_run.font.size = Pt(11)
desc_run.font.color.rgb = MUTED
desc_run.font.italic = True

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 1: BUILDER MODULE OVERVIEW
# ════════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Builder Module Overview", level=1)

add_body(doc,
    'The Builder Module is a feasibility intelligence layer within the Qnit Site Analysis Tool. '
    'When a user selects the "builder" profile, the platform runs 9 structured signal checks across '
    'zoning, planning regulations, terrain, connectivity, utilities, and market context -- returning a '
    'GO / CAUTION / NO-GO verdict with confidence tiers for each signal.',
    size=10)

add_body(doc, "Key characteristics:", bold=True, size=10)
bullets_overview = [
    "Profile trigger: useProfileStore({ profile: \"builder\" }) → shows Zoning + Title & Documents modules instead of climate modules",
    "Entry point: apps/web/app/project/new/page.tsx → runBuilderSignals() in apps/web/lib/api/verdict.ts",
    "9 core signals: zoneRing, farAssembly, obligations, overlays, terrain, connectivitySignal, utilities, priceUpside, growth",
    "Confidence tiers: authoritative → inferred → unresolved (never fabricated, absence ≠ clear)",
    "Verdict engine: POST /report/verdict aggregates signal scores → GO / CAUTION / NO-GO",
    "Premium FAR: computed client-side from UDD 78 MNJ 2024(E) gazette band rules (service implementation pending US-085)",
    "Supabase: projects persisted in sat_projects table (RLS: user_id = auth.uid())",
]
for b in bullets_overview:
    add_bullet(doc, b)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 2: FEATURE DATA SOURCES
# ════════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Feature Data Sources", level=1)
add_body(doc,
    "Each feature/signal in the builder module pulls from one or more external data sources. "
    "The table below maps every feature to its primary source, fallback, service, and confidence tier.",
    size=10)

FEATURES = [
    # (Feature, Signal ID, Service, Port, Primary Sources, Fallback, Confidence)
    ("Zone & Ring (RMP-2015)", "zoneRing", "geo", "8005",
     "BDA RMP-2015 via KGIS MapServer; OSM (inferred fallback); ISRO NRSC Bhuvan LULC",
     "OSM Overpass (Core Ring Road / ORR polygon)", "authoritative / inferred"),
    ("FAR Assembly", "farAssembly", "planning", "8006",
     "NBC 2016 Table 15 (base FAR brackets); BDA CDP 2031 (ring modifiers); BDA TOD Notification 2020 (FAR 4.0 metro proximity)",
     "OSM road-width tag; BMRCL metro presence", "inferred"),
    ("Obligations: Mixed-use / Parking / TIA", "obligations", "planning", "8006",
     "BDA RMP-2015 Ch.8 Table 23 (parking ECS); Reg 4.1.2 / 4.2.2 / 7.3 (mixed-use %) ; BBMP/BDA TIA trigger thresholds (curated)",
     "NBC 2016 access minimums", "inferred"),
    ("Deal-Killer Overlays", "overlays", "geo", "8005",
     "OSM Overpass (rajakaluve, HT power lines, lakes, wetlands); AAI ARP (bundled airport points); WRIS (water resources); ISRO Bhuvan LULC (forest); IGL/GAIL CGD mapping (gas pipelines)",
     "Absence ≠ clear (unresolved when no bundled clearing layer)", "unresolved"),
    ("Terrain / Slope / Geotech", "terrain", "flood", "8002",
     "Copernicus GLO-30 DEM (slope %, elevation); MERIT Hydro (HAND — Height Above Nearest Drainage); SRTM (elevation grid)",
     "Manual geotechnical survey (bearing capacity — authoritative only when supplied)", "inferred / authoritative"),
    ("Connectivity: Airport / Metro / Highway", "connectivitySignal", "infrastructure", "8007",
     "AAI ARP bundled dataset (authoritative airport locations); OSM Overpass (metro/rail/highway nearest features)",
     "BMRCL GTFS (metro confirmation, optional)", "inferred"),
    ("Utilities & NOC Checklist", "utilities", "infrastructure", "8007",
     "BWSSB trunk-main layers (authoritative, when available); OSM Overpass (water/telecom/power/drainage proximity proxy)",
     "BESCOM (power substation context); KSPCB, Fire, AAI NOCAS, PNGRB, Telecom RoW (NOC checklist items)", "unknown / inferred"),
    ("Price Upside (Indicative)", "priceUpside", "future-infra", "8008",
     "Kaveri Guidance Value (user-supplied ₹/sqm benchmark); Curated infrastructure pipeline (operational + under-construction)",
     "BDA / BMRCL / NHAI public announcements", "inferred"),
    ("Growth Pipeline", "growth", "future-infra", "8008",
     "Curated pipeline dataset — BMRCL, BDA, NHAI, KIADB (snapshot: 2024-Q4)",
     "Project status: Operational / Under Construction / Cancelled / Tendered", "inferred"),
    ("Verdict Engine (GO / CAUTION / NO-GO)", "verdict", "report", "8010",
     "Aggregated builder signal scores from all 9 panels",
     "—", "computed"),
    ("Gas Pipelines (map overlay)", "—", "cadastral", "8011",
     "IGL/GAIL City Gas Distribution (CGD) pipeline mapping",
     "—", "inferred"),
    ("BWSSB Sewerage (map overlay)", "—", "cadastral", "8011",
     "BWSSB sewerage network / data.opencity.in",
     "—", "inferred"),
    ("BBMP Storm Drains (map overlay)", "—", "cadastral", "8011",
     "BBMP storm water drain network / data.opencity.in",
     "—", "inferred"),
    ("Encroachment (map overlay)", "—", "cadastral", "8011",
     "Karnataka Revenue Department encroachment records",
     "—", "inferred"),
    ("Water Bodies / Drainage (map overlay)", "—", "cadastral", "8011",
     "WRIS (Water Resources Information System) lake and drainage polygons",
     "—", "inferred"),
    ("Power Grid (map overlay)", "—", "Overpass API", "external",
     "OpenStreetMap (ODbL) — 11kV+ BESCOM distribution; ≥66kV KPTCL transmission lines + substations",
     "—", "inferred"),
    ("Premium FAR", "—", "client-side", "n/a",
     "UDD Gazette No. 78 MNJ 2024(E) dated 21-Feb-2025 — road-width band rules: <9m=0%, 9–12m=+20%, 12–18m=+40%, ≥18m=+40%(+20% TDR unpriced)",
     "—", "indicative"),
    ("Amenities Nearby", "—", "geo", "8005",
     "OpenStreetMap (ODbL) via Overpass API — schools, hospitals, banks, supermarkets, parks within radius",
     "—", "inferred"),
]

headers = ["Feature", "Signal ID", "Svc", "Port", "Primary Data Sources", "Fallback / Notes", "Confidence"]
col_widths = [Cm(3.4), Cm(2.2), Cm(1.4), Cm(1.1), Cm(5.8), Cm(3.2), Cm(2.0)]

tbl = doc.add_table(rows=1 + len(FEATURES), cols=len(headers))
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_border(tbl)

# Header row
for ci, (h, w) in enumerate(zip(headers, col_widths)):
    c = tbl.cell(0, ci)
    c.width = w
    set_cell_bg(c, GREEN)
    cell_para(c, h, bold=True, color=WHITE, size=8.5)

# Data rows
for ri, row_data in enumerate(FEATURES):
    for ci, (val, w) in enumerate(zip(row_data, col_widths)):
        c = tbl.cell(ri + 1, ci)
        c.width = w
        bg = GREEN_BG if ri % 2 == 0 else CREAM
        set_cell_bg(c, bg)
        color = AMBER if ci == 6 and val == "unresolved" else (MUTED if ci == 6 else DARK)
        size = 8 if ci in (4, 5) else 8.5
        cell_para(c, val, color=color, size=size)

doc.add_paragraph()
add_body(doc,
    "Confidence definitions: authoritative = primary government/regulatory source; "
    "inferred = derived from secondary data (OSM, DEM, curated); "
    "unresolved = data unavailable or absence cannot confirm safety; "
    "indicative = computed from published gazette rules, not verified by authority.",
    color=MUTED, size=8.5, italic=True)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 3: API ENDPOINTS
# ════════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. API Endpoints Reference", level=1)
add_body(doc,
    "All builder module services run as FastAPI microservices. Base URLs are set via environment variables. "
    "Every endpoint is feature-flag gated — returns HTTP 403 unless the corresponding flag is in FLAGS= env var.",
    size=10)

SERVICES = [
    {
        "name": "Geo Service",
        "port": 8005,
        "env": "NEXT_PUBLIC_GEO_API_URL",
        "default": "http://localhost:8005",
        "endpoints": [
            ("GET", "/geo/zone-resolve", "lat, lon", "feature.geo.zone-resolver", "Zone class + authority tier (RMP → OSM → Bhuvan tiered resolver)"),
            ("GET", "/geo/ring", "lat, lon", "feature.zoning.land-use", "RMP-2015 planning ring classification (I/II/III = TDR Zone A/B/C)"),
            ("GET", "/geo/overlays", "lat, lon, radius_m", "feature.geo.overlays", "Unified deal-killer overlay engine — rajakaluve, airport OLS, lakes, wetlands, forest, HT lines"),
            ("GET", "/geo/amenities", "lat, lon, radius_m", "feature.geo.amenities", "Nearby amenities from OpenStreetMap Overpass"),
            ("GET", "/geo/transport-access", "lat, lon", "feature.geo.transport-access", "Nearest transport access nodes with coordinates"),
        ]
    },
    {
        "name": "Planning Service",
        "port": 8006,
        "env": "NEXT_PUBLIC_PLANNING_API_URL",
        "default": "http://localhost:8006",
        "endpoints": [
            ("POST", "/planning/far", "{ lat, lon, zone_class, plot_area_sqm }", "feature.planning.far-assembly", "Permissible vs achievable FAR assembly — NBC 2016 + BDA CDP 2031 + TOD"),
            ("POST", "/planning/obligations", "{ lat, lon, zone_class, far_achievable, plot_area_sqm }", "feature.planning.mixed-use", "Mixed-use %, parking ECS, TIA trigger (BDA RMP-2015 Ch.8)"),
        ]
    },
    {
        "name": "Flood / Terrain Service",
        "port": 8002,
        "env": "NEXT_PUBLIC_FLOOD_API_URL",
        "default": "http://localhost:8002",
        "endpoints": [
            ("POST", "/flood/terrain", "{ geometry: GeoJSON Polygon }", "feature.flood.terrain", "Slope %, HAND, cut-fill estimate — Copernicus GLO-30 DEM + MERIT Hydro"),
        ]
    },
    {
        "name": "Infrastructure Service",
        "port": 8007,
        "env": "NEXT_PUBLIC_INFRASTRUCTURE_API_URL",
        "default": "http://localhost:8007",
        "endpoints": [
            ("POST", "/infrastructure/connectivity", "{ lat, lon }", "feature.infrastructure.connectivity", "Airport / metro / rail / highway distances — AAI ARP + OSM Overpass"),
            ("POST", "/infrastructure/utilities", "{ lat, lon }", "feature.infrastructure.utilities", "Water main, power, sewage presence + NOC checklist — BWSSB + OSM"),
            ("GET", "/infrastructure/power-grid", "lat, lon, radius_m", "feature.infrastructure.power-grid", "KPTCL / BESCOM power line proximity — OpenStreetMap"),
        ]
    },
    {
        "name": "Future-Infra Service",
        "port": 8008,
        "env": "NEXT_PUBLIC_FUTURE_INFRA_API_URL",
        "default": "http://localhost:8008",
        "endpoints": [
            ("POST", "/future-infra/price-upside", "{ lat, lon, guidance_value_per_sqm }", "feature.context.growth-pipeline", "Indicative price upside range — Kaveri guidance value + curated pipeline"),
            ("GET", "/future-infra/pipeline", "lat, lon, radius_km=10", "feature.context.growth-pipeline", "Infrastructure pipeline projects within radius — BMRCL/BDA/NHAI/KIADB (2024-Q4)"),
        ]
    },
    {
        "name": "Report / Verdict Service",
        "port": 8010,
        "env": "NEXT_PUBLIC_REPORT_API_URL",
        "default": "http://localhost:8010",
        "endpoints": [
            ("POST", "/report/verdict", "{ signals: BuilderSignalBundle }", "feature.report.go-no-go", "GO / CAUTION / NO-GO verdict + per-signal scores — aggregated builder signals"),
        ]
    },
    {
        "name": "Cadastral Service",
        "port": 8011,
        "env": "NEXT_PUBLIC_CADASTRAL_API_URL",
        "default": "http://localhost:8011",
        "endpoints": [
            ("GET", "/gas-pipelines", "lat, lon, radius_m", "feature.cadastral.overlays", "IGL/GAIL gas pipeline proximity polygons"),
            ("GET", "/bwssb-sewerage", "lat, lon, radius_m", "feature.cadastral.overlays", "BWSSB sewerage network proximity"),
            ("GET", "/bbmp-swd", "lat, lon, radius_m", "feature.cadastral.overlays", "BBMP storm water drain proximity"),
            ("GET", "/encroachment", "lat, lon, radius_m", "feature.cadastral.overlays", "Karnataka Revenue Dept encroachment records"),
            ("GET", "/wris-lakes", "lat, lon, radius_m", "feature.cadastral.overlays", "WRIS lake and drainage polygons"),
            ("GET", "/cadastral/records", "lat, lon", "feature.cadastral.land-records", "e-Chawadi (Bhoomi) land records — survey number, owner, RTC/EC links"),
        ]
    },
]

for svc in SERVICES:
    add_heading(doc, f"{svc['name']}  (port {svc['port']})", level=2)
    add_body(doc, f"Env var: {svc['env']}  ·  Default: {svc['default']}", color=MUTED, size=9)

    ep_tbl = doc.add_table(rows=1 + len(svc["endpoints"]), cols=5)
    ep_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(ep_tbl)

    ep_headers = ["Method", "Path", "Params", "Flag", "Description"]
    ep_widths = [Cm(1.4), Cm(3.8), Cm(3.8), Cm(4.2), Cm(5.8)]
    for ci, (h, w) in enumerate(zip(ep_headers, ep_widths)):
        c = ep_tbl.cell(0, ci)
        c.width = w
        set_cell_bg(c, GREEN)
        cell_para(c, h, bold=True, color=WHITE, size=8)

    for ri, (method, path, params, flag, desc) in enumerate(svc["endpoints"]):
        method_color = GREEN if method == "GET" else AMBER
        for ci, (val, w, col) in enumerate(zip(
            [method, path, params, flag, desc],
            ep_widths,
            [method_color, DARK, MUTED, RGBColor(0x1A, 0x4A, 0x8A), DARK]
        )):
            c = ep_tbl.cell(ri + 1, ci)
            c.width = w
            set_cell_bg(c, GREEN_BG if ri % 2 == 0 else CREAM)
            cell_para(c, val, color=col, size=8, bold=(ci == 0))

    doc.add_paragraph()

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 4: SUPABASE SETUP
# ════════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Supabase Setup", level=1)
add_body(doc,
    "Supabase is used for authentication (all profiles) and project persistence (sat_projects table). "
    "The existing Qnit Supabase project hosts both. No separate Supabase project is needed.",
    size=10)

add_heading(doc, "4.1  Create the sat_projects Table", level=2)
add_body(doc, "Run the following SQL in the Supabase dashboard → SQL Editor:", size=10)

DDL = """create table if not exists public.sat_projects (
  id            uuid          primary key default gen_random_uuid(),
  user_id       uuid          not null references auth.users(id) on delete cascade,
  name          text          not null,
  location      text,
  status        text          not null default 'needs-review',
  boundary      jsonb,
  coordinates   text,
  area_sqm      double precision,
  modules_run   text[],
  overall_score integer,
  created_at    timestamptz   not null default now()
);

-- Enable Row Level Security
alter table public.sat_projects enable row level security;

-- Policy: users can only see/edit their own projects
create policy "users see own projects"
  on public.sat_projects
  for all
  using  (user_id = auth.uid())
  with check (user_id = auth.uid());"""

add_code(doc, DDL)

add_heading(doc, "4.2  Environment Variables", level=2)
add_body(doc, "Set in apps/web/.env.local (never commit):", size=10)
add_code(doc, "NEXT_PUBLIC_SUPABASE_URL=https://<project-ref>.supabase.co\nNEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY=sb_publishable_<key>")
add_body(doc,
    "Get both values from: Supabase dashboard → Project Settings → API. "
    "The env var is PUBLISHABLE_KEY, not ANON_KEY.",
    color=MUTED, size=9, italic=True)

add_heading(doc, "4.3  Auth URL Configuration", level=2)
add_body(doc, "In Supabase dashboard → Authentication → URL Configuration:", size=10)
bullets_auth = [
    "Site URL: https://qnit.site",
    "Redirect allow-list: https://qnit.site/**, http://localhost:3000/**",
]
for b in bullets_auth:
    add_bullet(doc, b)

add_heading(doc, "4.4  Code Integration", level=2)
code_files = [
    ("Client", "apps/web/lib/supabase/client.ts", "createClient(NEXT_PUBLIC_SUPABASE_URL, NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY)"),
    ("CRUD API", "apps/web/lib/api/projects.ts", "getProjects(), getProject(id), createProject() — all query sat_projects via Supabase client"),
    ("Auth store", "apps/web/lib/stores/auth.ts", "useAuthStore — wraps supabase.auth.getUser() / signOut()"),
]
ct = doc.add_table(rows=1 + len(code_files), cols=3)
ct.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_border(ct)
for ci, h in enumerate(["Layer", "File", "Notes"]):
    c = ct.cell(0, ci)
    set_cell_bg(c, GREEN)
    cell_para(c, h, bold=True, color=WHITE, size=9)
    c.width = [Cm(2.4), Cm(6.2), Cm(10.4)][ci]
for ri, (layer, fpath, notes) in enumerate(code_files):
    bg = GREEN_BG if ri % 2 == 0 else CREAM
    for ci, (val, col) in enumerate([(layer, GREEN), (fpath, DARK), (notes, MUTED)]):
        c = ct.cell(ri + 1, ci)
        set_cell_bg(c, bg)
        cell_para(c, val, color=col, size=8.5)
doc.add_paragraph()

add_heading(doc, "4.5  Security Notes", level=2)
sec_bullets = [
    "RLS enforces user_id = auth.uid() for ALL operations — no server-side check needed in API layer",
    "Service-role key NOT used by frontend — only publishable key in env",
    "Pending: enable leaked-password protection in Supabase dashboard (Phase 6 carry-forward)",
    "Pending: revoke SECURITY DEFINER on public.rls_auto_enable() if present (Supabase advisor)",
]
for b in sec_bullets:
    add_bullet(doc, b)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 5: FEATURE FLAGS
# ════════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Feature Flags", level=1)
add_body(doc,
    "All builder features are default-off, gated by feature flags in packages/flags/src/flags.py. "
    "Each service reads os.getenv(\"FLAGS\") directly — flags are NOT imported from the Python package "
    "(outside Docker build context). Enable via FLAGS= env var.",
    size=10)

add_code(doc, 'FLAGS="feature.geo.zone-resolver,feature.zoning.land-use,feature.geo.overlays,\\\n'
              'feature.planning.far-assembly,feature.planning.mixed-use,feature.flood.terrain,\\\n'
              'feature.infrastructure.connectivity,feature.infrastructure.utilities,\\\n'
              'feature.infrastructure.power-grid,feature.context.growth-pipeline,\\\n'
              'feature.report.go-no-go,feature.cadastral.overlays,feature.cadastral.land-records,\\\n'
              'feature.geo.amenities,feature.geo.transport-access"')

FLAGS_TABLE = [
    ("feature.geo.zone-resolver", "Zone + authority tier resolver"),
    ("feature.zoning.land-use", "RMP-2015 ring classification (I/II/III)"),
    ("feature.geo.overlays", "Unified deal-killer overlay engine"),
    ("feature.planning.far-assembly", "FAR assembly (NBC 2016 + BDA CDP 2031 + TOD)"),
    ("feature.planning.mixed-use", "Obligations: mixed-use %, parking ECS, TIA"),
    ("feature.flood.terrain", "Terrain, slope, HAND, geotech"),
    ("feature.infrastructure.connectivity", "Connectivity: airport / metro / highway distances"),
    ("feature.infrastructure.utilities", "Utilities presence + NOC checklist"),
    ("feature.infrastructure.power-grid", "Power grid proximity (KPTCL/BESCOM)"),
    ("feature.context.growth-pipeline", "Price upside + growth pipeline (2024-Q4)"),
    ("feature.report.go-no-go", "Verdict engine: GO / CAUTION / NO-GO"),
    ("feature.cadastral.overlays", "Gas / BWSSB / BBMP / encroachment / WRIS map overlays"),
    ("feature.cadastral.land-records", "e-Chawadi Bhoomi land records lookup"),
    ("feature.cadastral.explorer", "Cadastral hierarchy browser"),
    ("feature.geo.amenities", "Amenities nearby (OpenStreetMap)"),
    ("feature.geo.transport-access", "Nearest transport access nodes"),
]

ft = doc.add_table(rows=1 + len(FLAGS_TABLE), cols=2)
ft.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_border(ft)
for ci, h in enumerate(["Flag", "Enables"]):
    c = ft.cell(0, ci)
    set_cell_bg(c, GREEN)
    cell_para(c, h, bold=True, color=WHITE, size=9)
    c.width = [Cm(8.0), Cm(11.0)][ci]
for ri, (flag, desc) in enumerate(FLAGS_TABLE):
    bg = GREEN_BG if ri % 2 == 0 else CREAM
    c0 = ft.cell(ri + 1, 0)
    c1 = ft.cell(ri + 1, 1)
    set_cell_bg(c0, bg)
    set_cell_bg(c1, bg)
    cell_para(c0, flag, color=RGBColor(0x1A, 0x4A, 0x8A), size=8.5)
    cell_para(c1, desc, size=8.5)
doc.add_paragraph()

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 6: DEVELOPER HANDOFF
# ════════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Developer Handoff", level=1)

add_heading(doc, "6.1  Non-Negotiable Rules", level=2)
rules = [
    "Contract-first — update contracts/<service>.yaml + contracts/CHANGELOG.md BEFORE writing service code. CI gate fails PR otherwise.",
    "Flag-default-off — every new behavior gated by a FeatureFlag enum value. Enable via FLAGS= env var only after validation.",
    "One feature per PR — tooling/refactor exceptions rare. PRs touching contracts/ must update CHANGELOG.md (CI enforced).",
    "No direct push to main — branch + PR + 1 review + green CI.",
    "No secrets in tracked files — .env, .claude/mcp.json, .claude/settings.local.json are gitignored. Never paste tokens or keys.",
    "FVD before code — new feature requires docs/feature-validation/SAT-XX_*.md with acceptance criteria mapped to functions.",
]
for r in rules:
    add_bullet(doc, r)

add_heading(doc, "6.2  Local Dev Setup — Builder Module", level=2)
add_code(doc,
"""# 1. Start all builder services (each needs its own .venv with Python 3.12)
cd services/geo          && source .venv/Scripts/activate && uvicorn app.main:app --reload --port 8005 &
cd services/planning     && source .venv/Scripts/activate && uvicorn app.main:app --reload --port 8006 &
cd services/flood        && source .venv/Scripts/activate && uvicorn app.main:app --reload --port 8002 &
cd services/infrastructure && source .venv/Scripts/activate && uvicorn app.main:app --reload --port 8007 &
cd services/future-infra && source .venv/Scripts/activate && uvicorn app.main:app --reload --port 8008 &
cd services/report       && source .venv/Scripts/activate && uvicorn app.main:app --reload --port 8010 &
cd services/cadastral    && source .venv/Scripts/activate && uvicorn app.main:app --reload --port 8011 &

# 2. Set all builder flags
export FLAGS="feature.geo.zone-resolver,feature.zoning.land-use,feature.geo.overlays,feature.planning.far-assembly,feature.planning.mixed-use,feature.flood.terrain,feature.infrastructure.connectivity,feature.infrastructure.utilities,feature.infrastructure.power-grid,feature.context.growth-pipeline,feature.report.go-no-go,feature.cadastral.overlays,feature.cadastral.land-records,feature.geo.amenities,feature.geo.transport-access"

# 3. Start frontend
cd apps/web && npm run dev   # http://localhost:3000

# 4. Select "Builder" profile → New Analysis → drop a pin → verify 9 signal panels load""")

add_heading(doc, "6.3  Key File Map", level=2)

FILE_MAP = [
    ("Frontend", "apps/web/lib/api/verdict.ts", "runBuilderSignals() — all 9 signal fetchers"),
    ("Frontend", "apps/web/lib/api/projects.ts", "Supabase CRUD for sat_projects table"),
    ("Frontend", "apps/web/lib/api/cadastral_records.ts", "e-Chawadi land records fetcher"),
    ("Frontend", "apps/web/lib/stores/analysis.ts", "ModuleResult type, signal state management"),
    ("Frontend", "apps/web/components/map/SitePlanningCard.tsx", "FAR/setback popup + calcPremiumFar()"),
    ("Frontend", "apps/web/components/map/ParcelVerdictDock.tsx", "Builder verdict signal dock panel"),
    ("Frontend", "apps/web/components/map/MapContainer.tsx", "amenitiesCenter prop — fixes overlay fetch location"),
    ("Frontend", "apps/web/app/project/new/page.tsx", "Builder profile detection, suggestName(), pin flow"),
    ("Frontend", "apps/web/app/dashboard/page.tsx", "BUILDER_PIPS, builder badge, project list"),
    ("Service", "services/geo/ (:8005)", "Zone, ring, overlays, amenities, transport-access"),
    ("Service", "services/planning/ (:8006)", "FAR assembly, obligations"),
    ("Service", "services/flood/ (:8002)", "Terrain, slope, HAND, geotech"),
    ("Service", "services/infrastructure/ (:8007)", "Connectivity, utilities, power-grid"),
    ("Service", "services/future-infra/ (:8008)", "Price upside, growth pipeline"),
    ("Service", "services/report/ (:8010)", "Verdict engine (GO/CAUTION/NO-GO)"),
    ("Service", "services/cadastral/ (:8011)", "e-Chawadi overlays + land records"),
    ("Contract", "contracts/geo.yaml", "Geo service OpenAPI spec"),
    ("Contract", "contracts/planning.yaml", "Planning service OpenAPI spec"),
    ("Contract", "contracts/infrastructure.yaml", "Infrastructure service OpenAPI spec"),
    ("Contract", "contracts/future-infra.yaml", "Future-infra service OpenAPI spec"),
    ("Contract", "contracts/report.yaml", "Report/verdict service OpenAPI spec"),
    ("Contract", "contracts/cadastral.yaml", "Cadastral service OpenAPI spec"),
]

fmtbl = doc.add_table(rows=1 + len(FILE_MAP), cols=3)
fmtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_border(fmtbl)
for ci, h in enumerate(["Layer", "Path", "Purpose"]):
    c = fmtbl.cell(0, ci)
    set_cell_bg(c, GREEN)
    cell_para(c, h, bold=True, color=WHITE, size=9)
    c.width = [Cm(2.0), Cm(6.8), Cm(10.2)][ci]
for ri, (layer, path, purpose) in enumerate(FILE_MAP):
    bg = GREEN_BG if ri % 2 == 0 else CREAM
    layer_color = GREEN if layer == "Frontend" else (AMBER if layer == "Service" else MUTED)
    for ci, (val, col) in enumerate([(layer, layer_color), (path, DARK), (purpose, MUTED)]):
        c = fmtbl.cell(ri + 1, ci)
        set_cell_bg(c, bg)
        cell_para(c, val, color=col, size=8.5)
doc.add_paragraph()

add_heading(doc, "6.4  Pending Items (Next Sprint)", level=2)
pending = [
    "US-085: Full gazette transcription UDD 78 MNJ 2024(E) → planning service. Premium FAR currently client-side estimate only (calcPremiumFar in SitePlanningCard.tsx).",
    "Cadastral service (:8011) prod hosting: 418 MB village index parquet needs VM with data volume or Supabase PostGIS hosting plan.",
    "GH#53 / GH#55: Supabase sat_projects table creation pending (send DDL in §4.1 to Chirag). Frontend is already wired — will fail silently until table exists.",
    "Download Report feature: deferred to next sprint. Code removed from results page. Needs: fetchLayerSummaries + runBuilderSignals + generateBuilderReportHtml + downloadBlob wired back with UI placement.",
    "KGIS RMP-2015 authoritative layer: license acquisition pending (phase-0-kgis-verification.md). Zone resolver currently falls back to OSM-inferred.",
]
for p in pending:
    add_bullet(doc, p)

add_heading(doc, "6.5  Production Deployment", level=2)
prod_info = [
    ("Frontend", "https://qnit.site", "Vercel project: qnit-web, root: apps/web"),
    ("API Gateway", "https://api.qnit.site", "Caddy reverse proxy on AWS EC2 t3.medium, Mumbai (ap-south-1)"),
    ("Supabase", "Project: SAT (Tokyo ap-northeast-1)", "Auth + sat_projects table"),
    ("DNS", "GoDaddy", "qnit.site A→76.76.21.21 (Vercel); api.→65.1.245.213 (EC2)"),
    ("Contract version", "v2.32.0", "See contracts/CHANGELOG.md for full history"),
]
pdtbl = doc.add_table(rows=1 + len(prod_info), cols=3)
pdtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_border(pdtbl)
for ci, h in enumerate(["Component", "URL / Identifier", "Notes"]):
    c = pdtbl.cell(0, ci)
    set_cell_bg(c, GREEN)
    cell_para(c, h, bold=True, color=WHITE, size=9)
    c.width = [Cm(3.0), Cm(6.0), Cm(10.0)][ci]
for ri, (comp, url, notes) in enumerate(prod_info):
    bg = GREEN_BG if ri % 2 == 0 else CREAM
    for ci, (val, col) in enumerate([(comp, GREEN), (url, DARK), (notes, MUTED)]):
        c = pdtbl.cell(ri + 1, ci)
        set_cell_bg(c, bg)
        cell_para(c, val, color=col, size=8.5)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════════
doc.save(OUT_PATH)
print(f"Generated: {OUT_PATH}")
