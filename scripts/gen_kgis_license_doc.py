"""Generate docs/KGIS_License_Acquisition_Plan.docx — how to obtain a K-GIS / KSRSAC data license.

Team action doc: what to license, why a license is mandatory (verbatim KGIS terms),
whom to approach, the licensing mechanism, a ready-to-send request letter, step-by-step
next steps, timeline, and risks. Mirrors the house style of scripts/gen_deployment_doc.py
and scripts/gen_talkinglands_brief.py. Source: the research spike at
.claude/plans/i-need-to-understand-velvety-pearl.md.
Run: python scripts/gen_kgis_license_doc.py
"""

from __future__ import annotations

import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# ── palette (SAT brand) ──────────────────────────────────────────────────────
BRAND = RGBColor(0x2E, 0x7D, 0x6F)      # primary teal
INK = RGBColor(0x1C, 0x24, 0x20)        # near-black
MUTED = RGBColor(0x5A, 0x6A, 0x66)
MONO_BG = "EAF2F1"                       # secondary tint


def shade(cell, hex_fill: str) -> None:
    tc = cell._tc.get_or_add_tcPr()
    shd = tc.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tc.append(shd)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = BRAND
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pbdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pbdr.makeelement(qn("w:bottom"), {})
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "2E7D6F")
        pbdr.append(bottom)
        pPr.append(pbdr)
    else:
        run.font.size = Pt(12.5)
        run.font.color.rgb = INK
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)


def body(doc: Document, text: str, bold=False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(4)


def bullets(doc: Document, items, numbered=False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for it in items:
        p = doc.add_paragraph(style=style)
        r = p.add_run(it)
        r.font.size = Pt(10.5)
        r.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(1)


def quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(10)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED
    shade_paragraph(p, MONO_BG)


def shade_paragraph(p, hex_fill: str) -> None:
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


def table(doc: Document, headers, rows) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], "2E7D6F")
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if ri % 2 == 1:
                shade(cells[i], "F4F8F7")
            para = cells[i].paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── document ─────────────────────────────────────────────────────────────────
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

# title page
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = tp.add_run("Obtaining a K-GIS / KSRSAC Data License")
tr.bold = True
tr.font.size = Pt(24)
tr.font.color.rgb = BRAND
st = doc.add_paragraph()
sr = st.add_run("Process, contacts & action plan to license Karnataka cadastral + zonation data for SAT Builders View")
sr.font.size = Pt(12)
sr.font.color.rgb = MUTED

meta = doc.add_paragraph()
meta.paragraph_format.space_before = Pt(8)
for label, val in [
    ("Date", datetime.date.today().isoformat()),
    ("Prepared by", "Tanmay C J"),
    ("Purpose", "Lawful path to license K-GIS data for commercial use (Builders View)"),
    ("Status", "Research spike output — pre-build, pre-funding"),
    ("Audience", "SAT team / founders"),
]:
    rr = meta.add_run(f"{label}:  ")
    rr.bold = True
    rr.font.size = Pt(10)
    rr.font.color.rgb = INK
    vv = meta.add_run(f"{val}\n")
    vv.font.size = Pt(10)
    vv.font.color.rgb = MUTED

# 1. Executive summary
heading(doc, "1. Executive Summary")
body(doc, "SAT's planned Builders View needs Karnataka cadastral parcels (survey-number → parcel polygon) "
          "and land-use zonation. The Karnataka GIS platform (K-GIS), run by KSRSAC, already exposes exactly "
          "these via public Web API services and SAT can build against them during a research/dev spike now. "
          "However, K-GIS data is explicitly non-commercial: a formal data-sharing license / agreement from "
          "KSRSAC is mandatory before any commercial launch or funding round that monetises this data.")
body(doc, "There is no online self-serve form and no published price list — the process is contact-driven: "
          "a formal request to KSRSAC, a data-sharing agreement / MoU, and a committee-set fee. Land records "
          "(RTC, mutation, encumbrance) sit with separate departments (Revenue/Bhoomi, Stamps & "
          "Registration/Kaveri) and need their own agreements. This doc lists what to license, why, whom to "
          "approach, the mechanism, a ready-to-send request letter, and a step-by-step plan with timeline.")

# 2. What we need licensed
heading(doc, "2. What We Need Licensed")
body(doc, "Priority is the spatial backbone (Karnataka-first). Records are a separate, later track.")
table(doc, ["Data", "K-GIS service / source", "Use in SAT"], [
    ["Survey → parcel polygon", "geomForSurveyNum (Web API #7; DD/UTM)", "Core “type survey number → parcel on map” feature"],
    ["Coordinate → survey number", "getlocationdetails (already integrated)", "Reverse lookup / GPS context"],
    ["Land-use zonation", "Fetching Zonation Data (Web API #12)", "Data-derived zone_class for planning engine"],
    ["Admin hierarchy & village IDs", "Web API #1–#6", "Resolve survey → village id for geometry calls"],
    ["Background map layers", "K-GIS WMS / WFS", "Cadastral / thematic overlays"],
    ["Bulk cadastral (optional)", "Survey-Hissa layers (data-sharing)", "Own-data fallback / offline parcel DB"],
])

# 3. Why a license is mandatory
heading(doc, "3. Why a License Is Mandatory")
body(doc, "K-GIS is publicly reachable, but its own published terms restrict use. Verbatim, from the K-GIS FAQ "
          "and Web API pages:")
quote(doc, "“Under no circumstances data will be used for any commercial purposes by anyone.”")
quote(doc, "“The geospatial data available in KSRSAC Website shall not be used for any legal purpose.”")
quote(doc, "“…any data leakage / misuse by any individual / firm / by any Organization will be subjected "
           "to legal proceedings.”")
bullets(doc, [
    "Spike / development on K-GIS now = acceptable (non-commercial, internal validation).",
    "Commercial launch (paid product / monetised data) = requires a formal KSRSAC data-sharing license first.",
    "Product must be positioned as preliminary / indicative due-diligence, NOT legal title verification "
    "(the “not for any legal purpose” clause; also matches Dishaank's own “notional, not legally valid” "
    "disclaimer and 3–10 m GPS error).",
])

# 4. Whom to approach
heading(doc, "4. Whom to Approach")
body(doc, "Spatial data sits with KSRSAC/SSLR; records sit with Revenue and Stamps & Registration; master "
          "plans with the development authorities. Start with KSRSAC (primary).")
table(doc, ["Organisation", "What they hold", "Route / contact"], [
    ["KSRSAC (primary)",
     "K-GIS spatial data: survey/hissa geometry, zonation, WMS/WFS, satellite layers",
     "Email kgissupport@ksrsac.in + formal letter to the Director/CEO, KSRSAC, DPAR (e-Governance), "
     "GoK, Doddabettahalli, Bengaluru-560097"],
    ["SSLR (Survey, Settlement & Land Records)",
     "Original survey maps, Tippani/Phodi, Hissa creation history",
     "Commissioner, SSLR Department, Revenue Dept, GoK"],
    ["Revenue Dept / Bhoomi",
     "RTC / Pahani, mutation records",
     "Bhoomi Monitoring Cell, Revenue Department, GoK"],
    ["Stamps & Registration / Kaveri",
     "Encumbrance certificates, registered deeds, market value",
     "Inspector General of Registration (IGR), Dept of Stamps & Registration"],
    ["BDA / BMRDA",
     "Master Plan / RMP, land-use zoning, road & reservation lands",
     "Planning wing, Bangalore Development Authority / BMRDA"],
    ["Startup Karnataka / KDEM (lever)",
     "Govt-startup engagement, warm intros, credibility, possible policy benefits",
     "Karnataka Innovation & Technology Society (KITS) / Karnataka Digital Economy Mission"],
])
body(doc, "Accelerator lever: K-GIS is built for government departments. A sponsoring department, or "
          "engagement via Startup Karnataka / KDEM, can shorten access and add credibility to the request.")

# 5. Licensing mechanism
heading(doc, "5. Licensing Mechanism")
bullets(doc, [
    "Governing framework: KSRSAC data-sharing policy (aligned to NDSAP — National Data Sharing & "
    "Accessibility Policy).",
    "Instrument: a data-sharing agreement / MoU between the company and KSRSAC; non-government / commercial "
    "users sign and pay.",
    "Pricing: committee-set, per-layer / per-area, quoted on request — NOT published online. (“Data not in "
    "portal → contact KSRSAC to determine access terms.”)",
    "Access form: Web API keys + WMS/WFS endpoints, and/or bulk export (Shape/KML), per the agreement scope.",
])

# 6. Document bundle to prepare
heading(doc, "6. Document Bundle to Prepare")
bullets(doc, [
    "Company registration / incorporation documents (entity that will hold the license).",
    "Use-case note: what the product does, who uses it, how K-GIS data is displayed (1–2 pages).",
    "Scope sheet: exact layers/services (geomForSurveyNum, zonation, WMS/WFS), coverage (Karnataka), "
    "access mode (API vs bulk), refresh frequency, expected request volume.",
    "Undertaking: non-misuse, no redistribution of raw data, attribution, and “not for legal purpose” "
    "acknowledgement.",
    "Point of contact + authorised signatory details.",
])

# 7. Step-by-step next steps
heading(doc, "7. Step-by-Step Next Steps")
table(doc, ["#", "Action", "Owner", "Est. timing"], [
    ["1", "Finalise data scope (layers, area, API vs bulk, volume)", "Product lead", "Days 1–3"],
    ["2", "Assemble document bundle (§6)", "Founder + legal", "Week 1"],
    ["3", "Send formal request to KSRSAC (email + signed letter)", "Founder", "Week 1 (Day 0 of external clock)"],
    ["4", "Open Startup Karnataka / KDEM engagement for warm intro", "Founder", "Weeks 1–2 (parallel)"],
    ["5", "KSRSAC review + data-sharing committee → pricing quote", "KSRSAC", "Weeks 2–6"],
    ["6", "Negotiate + sign data-sharing agreement / MoU; pay fee", "Founder + legal", "Weeks 6–10"],
    ["7", "Receive API keys / WFS endpoints; integrate behind feature flag", "Engineering", "After signing"],
    ["8", "Separately initiate Bhoomi (Revenue) + Kaveri (S&R) record agreements", "Founder", "Parallel, 3–6 months"],
])

# 8. Draft request letter / email
heading(doc, "8. Draft Request Letter / Email (ready to send)")
body(doc, "Adapt names/figures, attach the bundle from §6, and send to the KSRSAC Director with "
          "kgissupport@ksrsac.in in copy.")
quote(doc, "To: The Director, KSRSAC, DPAR (e-Governance), Government of Karnataka, Doddabettahalli, "
           "Bengaluru-560097   |   Cc: kgissupport@ksrsac.in")
quote(doc, "Subject: Request for a K-GIS data-sharing agreement — Survey/Hissa geometry, zonation and "
           "WMS/WFS services (Karnataka)")
quote(doc, "Respected Sir/Madam,")
quote(doc, "We are [Company / Qnit], building a site-analysis product that helps architects and builders "
           "evaluate land parcels. We wish to license K-GIS spatial data for use within our application and "
           "request the applicable data-sharing agreement and terms.")
quote(doc, "Data requested (Karnataka coverage): (a) Geometric Polygon Area on Survey Number "
           "(geomForSurveyNum); (b) Fetching Zonation Data; (c) administrative-hierarchy / village-code "
           "services; (d) relevant WMS/WFS layers. Preferred access: Web API + WFS endpoints with periodic "
           "refresh; bulk export if required.")
quote(doc, "Purpose: preliminary site due-diligence and visualisation for our users. We acknowledge that the "
           "data is not to be used for any legal purpose, will not be redistributed as raw data, and will be "
           "attributed to KSRSAC / K-GIS.")
quote(doc, "Request: please share the applicable data-sharing policy, the agreement / MoU format, commercial-"
           "use terms, and the pricing for the above scope. We are prepared to execute the agreement and remit "
           "the applicable fees. Company registration, use-case note and scope sheet are enclosed.")
quote(doc, "Regards, [Name] · [Title] · [Company] · [Phone] · [Email]")

# 9. Timeline & milestones
heading(doc, "9. Timeline & Milestones")
body(doc, "Indicative only — government procurement timelines vary. Plan the commercial launch so it does "
          "not hard-block on this.")
table(doc, ["Phase", "Activity", "Indicative window"], [
    ["Prep", "Scope + document bundle", "Week 0–1"],
    ["Submit", "Formal request to KSRSAC + KDEM lever", "Week 1"],
    ["Review", "KSRSAC committee review + pricing quote", "Week 2–6"],
    ["Execute", "Negotiate, sign MoU, pay, receive access", "Week 6–10"],
    ["Integrate", "Wire API/WFS behind feature flag", "Week 10+"],
    ["Records", "Bhoomi + Kaveri agreements (separate)", "3–6 months, parallel"],
])

# 10. Risks & mitigations
heading(doc, "10. Risks & Mitigations")
table(doc, ["Risk", "Mitigation"], [
    ["Procurement delay drags on", "Continue non-commercial spike; phase launch so it doesn't hard-block; "
                                   "keep own-data digitisation fallback designed."],
    ["KSRSAC declines commercial licensing", "Get the commercial-use position in writing early; if barred, "
                                             "pivot to own-data aggregation/digitisation fallback."],
    ["Per-area pricing too high", "Scope Karnataka-only / Bengaluru-first; pursue startup-policy concession "
                                  "via KDEM / Startup Karnataka."],
    ["“Not for legal purpose” clause", "Position product as indicative due-diligence; add disclaimers; do "
                                         "not market as legal title verification."],
    ["Records (RTC/EC) have no API", "v1 = deep-link + user upload; live data only after Bhoomi/Kaveri "
                                     "agreements."],
])

# 11. National context
heading(doc, "11. National Context (supports the ask)")
bullets(doc, [
    "India Geospatial Guidelines 2021 (DST) liberalised acquisition, use and sharing of geospatial data for "
    "Indian entities — no prior approval for most data; favourable backdrop.",
    "National Geospatial Policy 2022 pushes open, standardised geospatial data and a national framework.",
    "DILRMP / ULPIN: Karnataka has ~1.1 crore 14-digit ULPINs (lat/long-based parcel IDs) — the emerging "
    "single source of truth; adopt ULPIN as the parcel key where available.",
    "Caveat: cadastral / survey-of-record and land records remain under state revenue & survey departments — "
    "the state agreements above are still required.",
])

# 12. Sources
heading(doc, "12. Sources")
bullets(doc, [
    "K-GIS Web API services — kgis.ksrsac.in/kgis/webapi.aspx",
    "K-GIS FAQ (terms of use) — kgis.ksrsac.in/kgis/faqs.aspx",
    "KSRSAC — ksrsac.karnataka.gov.in",
    "DILRMP / ULPIN — dilrmp.gov.in, dolr.gov.in/ulpin",
    "India Geospatial Guidelines 2021 (DST); National Geospatial Policy 2022",
])

out = Path(r"C:\Users\tanny\OneDrive\Desktop\Site\SAT\docs\KGIS_License_Acquisition_Plan.docx")
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(out))
print("WROTE", out)
